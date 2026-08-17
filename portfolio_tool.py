import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
import io
import json
import firebase_admin
from firebase_admin import credentials, firestore
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
import string
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import re
import hashlib

# ============================================================
# SECTION: STREAMLIT PAGE CONFIG
# ============================================================
# This must be the first Streamlit command. Keeping it before secrets,
# session-state, and Firebase initialization prevents a missing/invalid
# secret from stopping the login page from rendering.
st.set_page_config(page_title="Chew Advisory - Portfolio Analyzer", layout="wide", page_icon="")
hide_streamlit_style = """
<style>
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
header {visibility: hidden;}
.stDeployButton {display: none !important;}
.stToolbar {display: none !important;}
[data-testid="stHeader"] {display: none !important;}
</style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# ============================================================
# SECTION: FIREBASE & EMAIL CONFIGURATION
# ============================================================
# Firebase is optional for rendering the app. If its service-account secret
# is absent or invalid, the login page still loads and Firestore usage
# tracking is disabled gracefully.
firebase_init_error = None
if "firebase_initialized" not in st.session_state:
    try:
        if firebase_admin._apps:
            st.session_state.firebase_initialized = True
        else:
            cred_dict = st.secrets.get("FIREBASE_SERVICE_ACCOUNT", {})
            if isinstance(cred_dict, str):
                cred_dict = json.loads(cred_dict)
            if not cred_dict:
                raise ValueError("FIREBASE_SERVICE_ACCOUNT is not configured")
            cred = credentials.Certificate(cred_dict)
            firebase_admin.initialize_app(cred)
            st.session_state.firebase_initialized = True
    except Exception as e:
        firebase_init_error = str(e)
        st.session_state.firebase_initialized = False
        st.session_state.firebase_init_error = firebase_init_error
else:
    firebase_init_error = st.session_state.get("firebase_init_error")

try:
    db = firestore.client() if firebase_admin._apps else None
except Exception:
    db = None

try:
    ADMIN_EMAIL = st.secrets.get("ADMIN_EMAIL", "cktchew@gmail.com")
    GMAIL_ADDRESS = st.secrets.get("GMAIL_ADDRESS", "cktchew@gmail.com")
    GMAIL_APP_PASSWORD = st.secrets.get("GMAIL_APP_PASSWORD", "")
except Exception:
    ADMIN_EMAIL = "cktchew@gmail.com"
    GMAIL_ADDRESS = "cktchew@gmail.com"
    GMAIL_APP_PASSWORD = ""

# ============================================================
# SECTION: SESSION STATE INITIALIZATION
# ============================================================
for key in ['authenticated', 'user_email', 'user_name', 'page', 'portfolio_data', 'otp_code', 'otp_email', 'show_otp_input', 'funds_df']:
    if key not in st.session_state:
        st.session_state[key] = False if key == 'authenticated' else (None if key in ['user_email', 'user_name', 'otp_code', 'otp_email'] else 'home' if key == 'page' else None if key == 'funds_df' else {})

# ============================================================
# SECTION: HELPER FUNCTIONS
# ============================================================
def safe_float(v):
    try:
        return float(str(v).replace(',', '').strip())
    except:
        return np.nan

def generate_otp():
    return ''.join(random.choices(string.digits, k=6))

def detect_year_columns(df):
    year_returns = []
    year_benchmarks = []
    for col in df.columns:
        if re.match(r'^\d{4}\s+Return\s+\(%\)$', col):
            year = int(col.split()[0])
            year_returns.append((year, col))
        elif re.match(r'^\d{4}\s+Benchmark\s+\(%\)$', col):
            year = int(col.split()[0])
            year_benchmarks.append((year, col))
    year_returns.sort(key=lambda x: x[0])
    year_benchmarks.sort(key=lambda x: x[0])
    return year_returns, year_benchmarks

import io
import os
import re
import tempfile
from pathlib import Path
import numpy as np
import pandas as pd
from pypdf import PdfReader
from pdfminer.high_level import extract_text


def annualized_from_cumulative(value, years):
    if value is None or pd_nan(value):
        return np.nan
    value = float(value)
    if value <= -100:
        return np.nan
    return ((1.0 + value / 100.0) ** (1.0 / years) - 1.0) * 100.0


def cumulative_from_annualized(value, years):
    if value is None or pd_nan(value):
        return np.nan
    return ((1.0 + float(value) / 100.0) ** years - 1.0) * 100.0


def pd_nan(value):
    try:
        return bool(np.isnan(value))
    except Exception:
        return False


def _clean_text(text):
    text = text.replace("\u2014", "-").replace("\u2013", "-")
    text = text.replace("\u00a0", " ")
    return text


def _lines(text):
    return [re.sub(r"\s+", " ", line).strip() for line in _clean_text(text).splitlines() if line.strip()]


def _num_tokens(value):
    if not value:
        return []
    value = value.replace(",", "")
    matches = re.findall(r"[-+]?(?:\d+(?:\.\d*)?|\.\d+)", value)
    return [float(x) for x in matches]


def _numeric_only_line(line):
    line = line.strip().replace("%", "").replace(",", "")
    if not line:
        return []
    parts = line.split()
    values = []
    for part in parts:
        part = part.rstrip(".;,:")
        if part in {"-", "--", "N/A", "NA"}:
            values.append(np.nan)
        elif re.fullmatch(r"[-+]?(?:\d+(?:\.\d*)?|\.\d+)", part):
            values.append(float(part))
        else:
            return []
    return values


def _marker_index(lines, pattern, start=0):
    regex = re.compile(pattern, re.I)
    for i in range(start, len(lines)):
        if regex.search(lines[i]):
            return i
    return None


def _row_values(lines, pattern, occurrence=1, count=2, start=0, stop=None, lookahead=8):
    regex = re.compile(pattern, re.I)
    seen = 0
    stop = len(lines) if stop is None else min(stop, len(lines))
    for i in range(start, stop):
        match = regex.search(lines[i])
        if not match:
            continue
        seen += 1
        if seen != occurrence:
            continue
        values = _num_tokens(lines[i][match.end():])
        j = i + 1
        while len(values) < count and j < stop and j <= i + lookahead:
            next_values = _numeric_only_line(lines[j])
            if next_values:
                values.extend(next_values)
            j += 1
        return values[:count]
    return []


def _pairs_after(lines, start_pattern, count_pairs, stop_pattern=None):
    start = _marker_index(lines, start_pattern)
    if start is None:
        return []
    stop = len(lines)
    if stop_pattern:
        found_stop = _marker_index(lines, stop_pattern, start + 1)
        if found_stop is not None:
            stop = found_stop
    values = []
    for line in lines[start + 1:stop]:
        line_values = _numeric_only_line(line)
        if line_values:
            values.extend(line_values)
        if len(values) >= count_pairs * 2:
            break
    return [(values[i], values[i + 1]) for i in range(0, count_pairs * 2, 2)] if len(values) >= count_pairs * 2 else []


def _year_rows(lines, start=0, class_index=0, benchmark_index=1, max_years=30):
    result = {}
    year_re = re.compile(r"^(20\d{2})\.?\s*(.*)$")
    for i in range(start, len(lines)):
        match = year_re.match(lines[i])
        if not match:
            continue
        year = int(match.group(1))
        if year < 1990 or year > 2100:
            continue
        values = _num_tokens(match.group(2))
        j = i + 1
        while len(values) <= max(class_index, benchmark_index) and j < len(lines):
            if year_re.match(lines[j]):
                break
            values.extend(_numeric_only_line(lines[j]))
            j += 1
        if len(values) > max(class_index, benchmark_index):
            result[year] = (values[class_index], values[benchmark_index])
    return result


def _label_year_row(lines, pattern, start=0, count=5):
    values = _row_values(lines, pattern, occurrence=1, count=count, start=start, lookahead=2)
    return values


def _first_float(patterns, text):
    for pattern in patterns:
        match = re.search(pattern, text, re.I | re.S)
        if match:
            try:
                return float(match.group(1))
            except Exception:
                pass
    return np.nan


def _metadata(text, provider):
    compact = re.sub(r"\s+", " ", text)
    volatility = _first_float([
        r"Volatility Factor.*?\bis\s+(\d+(?:\.\d+)?)",
        r"volatility.{0,100}?(\d+(?:\.\d+)?)",
    ], compact)
    if provider == "TA":
        volatility = np.nan
    fee = _first_float([
        r"Annual Management Fee.{0,100}?(\d+(?:\.\d+)?)\s*%",
        r"Management Fee.{0,100}?(\d+(?:\.\d+)?)\s*%",
    ], compact)
    if provider == "AHAM" and not np.isfinite(fee):
        fee = _first_float([r"Max\s+(\d+(?:\.\d+)?)\s*%\s*per annum"], compact)

    date = ""
    date_match = re.search(r"as (?:at|of)\s+(\d{1,2}[ /-][A-Za-z0-9]+[ /-]\d{2,4})", compact, re.I)
    if date_match:
        date = date_match.group(1)
    else:
        date_match = re.search(r"as of date:\s*(\d{1,2}/\d{1,2}/\d{4})", compact, re.I)
        if date_match:
            date = date_match.group(1)

    benchmark_patterns = [
        r"MSCI AC Asia Pacific ex Japan High Dividend Yield Index",
        r"BPAM Corporates All Bond Index",
        r"MSCI Emerging Markets Net Total Return Index",
        r"Dow Jones Islamic Market World Index",
        r"MSCI World / Health Care NR USD Index",
        r"42% S&P 500 \+ 36% MSCI Europe \+ 12% MSCI Japan \+ 10% CIMB Bank 1-Month Fixed Deposit Rate",
        r"MSCI ACWI Information Technology Index \+ MSCI ACWI Communication Services Index",
    ]
    benchmark = ""
    for pattern in benchmark_patterns:
        match = re.search(pattern, compact, re.I)
        if match:
            benchmark = match.group(0)
            break
    return volatility, fee, date, benchmark


def _base_record(filename, text, provider, fund_name, share_class):
    volatility, fee, as_of_date, benchmark = _metadata(text, provider)
    return {
        "Fund Name": fund_name,
        "Provider": provider,
        "Share Class": share_class,
        "As Of Date": as_of_date,
        "1Y Return (%)": np.nan,
        "3Y Return (%)": np.nan,
        "5Y Return (%)": np.nan,
        "3Y Cumulative (%)": np.nan,
        "3Y Annualised (%)": np.nan,
        "5Y Cumulative (%)": np.nan,
        "5Y Annualised (%)": np.nan,
        "3Y Benchmark Cumulative (%)": np.nan,
        "3Y Benchmark Annualised (%)": np.nan,
        "5Y Benchmark Cumulative (%)": np.nan,
        "5Y Benchmark Annualised (%)": np.nan,
        "Volatility (%)": volatility,
        "Mgmt Fee (%)": fee,
        "1Y Benchmark (%)": np.nan,
        "Benchmark Name": benchmark,
        "Return Basis": "",
        "Source File": filename,
        "Extraction Warnings": "",
        "Calendar": {},
        "YTD": (np.nan, np.nan),
    }


def _set_period(record, period, fund_cum=np.nan, fund_ann=np.nan, bench_cum=np.nan, bench_ann=np.nan, basis=""):
    years = int(period.rstrip("Y"))
    if not np.isfinite(fund_cum) and np.isfinite(fund_ann):
        fund_cum = cumulative_from_annualized(fund_ann, years)
    if not np.isfinite(fund_ann) and np.isfinite(fund_cum):
        fund_ann = annualized_from_cumulative(fund_cum, years)
    if not np.isfinite(bench_cum) and np.isfinite(bench_ann):
        bench_cum = cumulative_from_annualized(bench_ann, years)
    if not np.isfinite(bench_ann) and np.isfinite(bench_cum):
        bench_ann = annualized_from_cumulative(bench_cum, years)

    record[f"{period} Cumulative (%)"] = fund_cum
    record[f"{period} Annualised (%)"] = fund_ann
    record[f"{period} Benchmark Cumulative (%)"] = bench_cum
    record[f"{period} Benchmark Annualised (%)"] = bench_ann
    record[f"{period} Return (%)"] = fund_ann
    record[f"{period} Benchmark (%)"] = bench_ann
    record[f"{period} Return Basis"] = basis


def _append_warning(record, warning):
    current = record.get("Extraction Warnings", "")
    record["Extraction Warnings"] = "; ".join([x for x in [current, warning] if x])


def _apply_calendar(record, calendar, ytd=(np.nan, np.nan)):
    record["Calendar"] = calendar or {}
    record["YTD"] = ytd


def _parse_aham(lines, record):
    fund_pattern = r"^Fund \(MYR\)"
    benchmark_pattern = r"^Benchmark \(MYR\)"
    total_start = _marker_index(lines, r"Total Return \(%\)")
    annual_start = _marker_index(lines, r"Annualised Return \(%\)")
    calendar_start = _marker_index(lines, r"Calendar Year Return \(%\)")

    fund_total = _row_values(lines, fund_pattern, count=4, start=total_start or 0) if total_start is not None else []
    bench_total = _row_values(lines, benchmark_pattern, count=4, start=total_start or 0) if total_start is not None else []
    fund_annual = _row_values(lines, fund_pattern, count=4, start=annual_start or 0) if annual_start is not None else []
    bench_annual = _row_values(lines, benchmark_pattern, count=4, start=annual_start or 0) if annual_start is not None else []

    if len(fund_total) >= 3:
        record["1Y Return (%)"] = fund_total[1]
        _set_period(record, "3Y", fund_cum=fund_total[2], bench_cum=bench_total[2] if len(bench_total) >= 3 else np.nan,
                     basis="Cumulative 3Y from FFS; annualised equivalent calculated")
    if len(fund_annual) >= 3:
        if not np.isfinite(record["1Y Return (%)"]):
            record["1Y Return (%)"] = fund_annual[0]
        _set_period(record, "3Y", fund_ann=fund_annual[1], bench_ann=bench_annual[1] if len(bench_annual) >= 2 else np.nan,
                     fund_cum=record.get("3Y Cumulative (%)", np.nan), bench_cum=record.get("3Y Benchmark Cumulative (%)", np.nan),
                     basis="Annualised 3Y from FFS")
        _set_period(record, "5Y", fund_ann=fund_annual[2], bench_ann=bench_annual[2] if len(bench_annual) >= 3 else np.nan,
                     basis="Annualised 5Y from FFS")
    if annual_start is None and total_start is None:
        _append_warning(record, "Annualised/total return table was not detected")

    if calendar_start is not None:
        calendar = _year_rows(lines, start=calendar_start + 1, class_index=0, benchmark_index=1)
        fund_ytd = _row_values(lines, fund_pattern, count=1, start=calendar_start)
        bench_ytd = _row_values(lines, benchmark_pattern, count=1, start=calendar_start)
        _apply_calendar(record, calendar, (fund_ytd[0] if fund_ytd else np.nan, bench_ytd[0] if bench_ytd else np.nan))
    return record


def _parse_amfunds(lines, record):
    pairs = _pairs_after(lines, r"Cumulative Return \(%\)", 6, r"Annualised Return \(%\)")
    if pairs:
        record["1Y Return (%)"] = pairs[3][0]
        _set_period(record, "3Y", fund_cum=pairs[4][0], bench_cum=pairs[4][1], basis="Cumulative 3Y from FFS; annualised equivalent calculated")
        _set_period(record, "5Y", fund_cum=pairs[5][0], bench_cum=pairs[5][1], basis="Cumulative 5Y from FFS; annualised equivalent calculated")
        record["1Y Benchmark (%)"] = pairs[3][1]
        record["YTD"] = pairs[0]
    else:
        _append_warning(record, "Cumulative return table was not detected")
    calendar_start = _marker_index(lines, r"Calendar Year Return \(%\)")
    if calendar_start is not None:
        _apply_calendar(record, _year_rows(lines, start=calendar_start + 1), record.get("YTD", (np.nan, np.nan)))
    _append_warning(record, "PDF extraction was marked restricted; extraction proceeded in-app")
    return record


def _parse_eastspring(lines, record):
    pairs = _pairs_after(lines, r"PERFORMANCE TABLE", 8, r"OTHER INFORMATION")
    if pairs:
        record["1Y Return (%)"] = pairs[3][0]
        record["1Y Benchmark (%)"] = pairs[3][1]
        _set_period(record, "3Y", fund_cum=pairs[4][0], bench_cum=pairs[4][1], basis="Cumulative 3Y from FFS; annualised equivalent calculated")
        _set_period(record, "5Y", fund_cum=pairs[5][0], bench_cum=pairs[5][1], basis="Cumulative 5Y from FFS; annualised equivalent calculated")
        record["YTD"] = pairs[2]
    else:
        _append_warning(record, "Performance table was not detected")
    annual_start = _marker_index(lines, r"Annual Fund Performance")
    if annual_start is not None:
        _apply_calendar(record, _year_rows(lines, start=annual_start + 1), record.get("YTD", (np.nan, np.nan)))
    return record


def _parse_manulife(lines, record):
    pairs = _pairs_after(lines, r"Total return over the following periods", 7, r"Calendar year returns")
    if pairs:
        record["1Y Return (%)"] = pairs[3][0]
        record["1Y Benchmark (%)"] = pairs[3][1]
        _set_period(record, "3Y", fund_cum=pairs[4][0], bench_cum=pairs[4][1], basis="Cumulative 3Y from FFS; annualised equivalent calculated")
        _set_period(record, "5Y", fund_cum=pairs[5][0], bench_cum=pairs[5][1], basis="Cumulative 5Y from FFS; annualised equivalent calculated")
        record["YTD"] = pairs[2]
    else:
        _append_warning(record, "Total-return table was not detected")
    calendar_start = _marker_index(lines, r"Calendar year returns")
    if calendar_start is not None:
        _apply_calendar(record, _year_rows(lines, start=calendar_start + 1), record.get("YTD", (np.nan, np.nan)))
    return record


def _parse_principal(lines, record):
    pairs = _pairs_after(lines, r"Cumulative Performance \(%\)", 8, r"Calendar Year Returns")
    if pairs:
        record["1Y Return (%)"] = pairs[4][0]
        record["1Y Benchmark (%)"] = pairs[4][1]
        _set_period(record, "3Y", fund_cum=pairs[5][0], bench_cum=pairs[5][1], basis="Cumulative 3Y from FFS; annualised equivalent calculated")
        _set_period(record, "5Y", fund_cum=pairs[6][0], bench_cum=pairs[6][1], basis="Cumulative 5Y from FFS; annualised equivalent calculated")
        record["YTD"] = pairs[0]
    else:
        _append_warning(record, "Cumulative performance table was not detected")
    calendar_start = _marker_index(lines, r"Calendar Year Returns")
    if calendar_start is not None:
        _apply_calendar(record, _year_rows(lines, start=calendar_start + 1), record.get("YTD", (np.nan, np.nan)))
    return record


def _parse_maybank(lines, record):
    def row(label, occurrence=1, count=5):
        return _row_values(lines, rf"^{label}$", occurrence=occurrence, count=count)

    ytd = row("YTD", 1, 5)
    one = row("1Y", 1, 5)
    ann3 = row("3Y", 1, 5)
    ann5 = row("5Y", 1, 5)
    cum3 = row("3Y", 2, 5)
    cum5 = row("5Y", 2, 5)
    if len(one) >= 5:
        record["1Y Return (%)"] = one[2]
        record["1Y Benchmark (%)"] = one[4]
    if len(ann3) >= 5:
        record["3Y Annualised (%)"] = ann3[2]
        record["3Y Benchmark Annualised (%)"] = ann3[4]
    if len(ann5) >= 5:
        record["5Y Annualised (%)"] = ann5[2]
        record["5Y Benchmark Annualised (%)"] = ann5[4]
    if len(cum3) >= 5:
        record["3Y Cumulative (%)"] = cum3[2]
        record["3Y Benchmark Cumulative (%)"] = cum3[4]
    if len(cum5) >= 5:
        record["5Y Cumulative (%)"] = cum5[2]
        record["5Y Benchmark Cumulative (%)"] = cum5[4]
    _set_period(record, "3Y", fund_cum=record["3Y Cumulative (%)"], fund_ann=record["3Y Annualised (%)"],
                 bench_cum=record["3Y Benchmark Cumulative (%)"], bench_ann=record["3Y Benchmark Annualised (%)"], basis="Both cumulative and annualised 3Y values from FFS")
    _set_period(record, "5Y", fund_cum=record["5Y Cumulative (%)"], fund_ann=record["5Y Annualised (%)"],
                 bench_cum=record["5Y Benchmark Cumulative (%)"], bench_ann=record["5Y Benchmark Annualised (%)"], basis="Both cumulative and annualised 5Y values from FFS")
    record["YTD"] = (ytd[2], ytd[4]) if len(ytd) >= 5 else (np.nan, np.nan)
    calendar_start = _marker_index(lines, r"Calendar Year Return")
    if calendar_start is not None:
        _apply_calendar(record, _year_rows(lines, start=calendar_start + 1, class_index=2, benchmark_index=4), record.get("YTD", (np.nan, np.nan)))
    if not len(ann3) >= 5 or not len(ann5) >= 5:
        _append_warning(record, "One or more annualised Maybank rows require manual review")
    return record


def _parse_ta(lines, record):
    cumulative_start = _marker_index(lines, r"Cumulative Fund Performance") or 0
    fund = _row_values(lines, r"^TA Global Technology Fund MYR Hdg", count=8, start=cumulative_start)
    benchmark = _row_values(lines, r"^TAGTF Benchmark", count=8, start=cumulative_start)
    if len(fund) >= 6:
        record["1Y Return (%)"] = fund[2]
        record["3Y Cumulative (%)"] = fund[3]
        record["5Y Cumulative (%)"] = fund[4]
        record["1Y Benchmark (%)"] = benchmark[2] if len(benchmark) >= 3 else np.nan
        _set_period(record, "3Y", fund_cum=fund[3], bench_cum=benchmark[3] if len(benchmark) >= 4 else np.nan,
                     basis="Cumulative 3Y from FFS; annualised equivalent calculated")
        _set_period(record, "5Y", fund_cum=fund[4], bench_cum=benchmark[4] if len(benchmark) >= 5 else np.nan,
                     basis="Cumulative 5Y from FFS; annualised equivalent calculated")
        record["YTD"] = (fund[5], benchmark[5] if len(benchmark) >= 6 else np.nan)
    else:
        _append_warning(record, "TA cumulative performance row was not detected")
    annual_start = _marker_index(lines, r"Annual Fund Performance")
    if annual_start is not None:
        annual_fund = _row_values(lines, r"^TA Global Technology Fund MYR Hdg", count=5, start=annual_start)
        annual_bench = _row_values(lines, r"^TAGTF Benchmark", count=5, start=annual_start)
        years = [2025, 2024, 2023, 2022, 2021]
        calendar = {}
        for year, value in zip(years, annual_fund):
            calendar[year] = (value, np.nan)
        for year, value in zip(years, annual_bench):
            calendar[year] = (calendar.get(year, (np.nan, np.nan))[0], value)
        _apply_calendar(record, calendar, record.get("YTD", (np.nan, np.nan)))
    _append_warning(record, "TA volatility badge is not exposed cleanly in the text layer; verify manually")
    return record


def _identify_provider(filename, text):
    low = f"{filename} {text}".lower()
    if "aham" in low:
        return "AHAM"
    if "maybank global sustainable" in low or "mgseif" in low:
        return "Maybank"
    if "eastspring" in low and "emerging markets" in low:
        return "Eastspring"
    if "manulife" in low and "healthcare" in low:
        return "Manulife"
    if "principal" in low and "global titans" in low:
        return "Principal"
    if "ta global technology" in low or "tagtf" in low:
        return "TA"
    if "ambond" in low or "amfunds" in low or "bpam" in low:
        return "AmFunds"
    return "Unknown"


def _identity(filename, text, provider):
    low = f"{filename} {text}".lower()
    if provider == "AHAM" and "asia pacific" in low:
        return "AHAM Select Asia Pacific (ex Japan) Dividend Fund (MYR Class)", "MYR"
    if provider == "AHAM":
        return "AHAM Select Bond Fund (MYR Class)", "MYR"
    if provider == "AmFunds":
        return "AmBond", "MYR"
    if provider == "Eastspring":
        return "Eastspring Investments Global Emerging Markets Fund", "MYR"
    if provider == "Manulife":
        return "Manulife Global Healthcare Fund (Fund A RM Hedged Class)", "Fund A (RM Hedged)"
    if provider == "Maybank":
        return "Maybank Global Sustainable Equity-I Fund (MYRH Class)", "MYRH"
    if provider == "Principal":
        return "Principal Global Titans Fund (Class MYR)", "MYR"
    if provider == "TA":
        return "TA Global Technology Fund (MYR Hedged Class)", "MYR Hedged"
    return Path(filename).stem, ""


def parse_ffs_text(text, filename):
    text = _clean_text(text)
    lines = _lines(text)
    provider = _identify_provider(filename, text)
    fund_name, share_class = _identity(filename, text, provider)
    record = _base_record(filename, text, provider, fund_name, share_class)
    if provider == "AHAM":
        record = _parse_aham(lines, record)
    elif provider == "AmFunds":
        record = _parse_amfunds(lines, record)
    elif provider == "Eastspring":
        record = _parse_eastspring(lines, record)
    elif provider == "Manulife":
        record = _parse_manulife(lines, record)
    elif provider == "Maybank":
        record = _parse_maybank(lines, record)
    elif provider == "Principal":
        record = _parse_principal(lines, record)
    elif provider == "TA":
        record = _parse_ta(lines, record)
    else:
        _append_warning(record, "Unknown FFS template; manual extraction review required")
    return record


def parse_ffs_bytes(pdf_bytes, filename):
    # Prefer pypdf's standard extraction. It preserves usable row text for the
    # sample FFSs, including the rotated TA factsheet. pdfminer is retained as
    # an in-process fallback for PDFs where pypdf returns little or no text.
    text = ""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes), strict=False)
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:
        text = ""

    if len(re.sub(r"\s+", "", text)) < 80:
        temp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as handle:
                handle.write(pdf_bytes)
                temp_path = handle.name
            text = extract_text(temp_path)
        finally:
            if temp_path and os.path.exists(temp_path):
                os.unlink(temp_path)

    if not text or len(re.sub(r"\s+", "", text)) < 80:
        raise ValueError("No usable selectable text was found. This appears to be an image-only PDF and cannot be processed under the no-OCR policy.")
    return parse_ffs_text(text, filename)


def records_to_dataframe(records):
    rows = []
    years = sorted({year for record in records for year in record.get("Calendar", {})})
    ytd_years = sorted({int(str(record.get("As Of Date", ""))[-4:]) for record in records if record.get("YTD") and any(np.isfinite(x) for x in record.get("YTD", (np.nan, np.nan)))})
    for record in records:
        row = {key: value for key, value in record.items() if key not in {"Calendar", "YTD"}}
        for year in years:
            fund, benchmark = record.get("Calendar", {}).get(year, (np.nan, np.nan))
            row[f"{year} Return (%)"] = fund
            row[f"{year} Benchmark (%)"] = benchmark
        as_of_year = None
        match = re.search(r"(20\d{2})", str(record.get("As Of Date", "")))
        if match:
            as_of_year = int(match.group(1))
        fund_ytd, benchmark_ytd = record.get("YTD", (np.nan, np.nan))
        for year in ytd_years:
            row[f"{year} YTD Return (%)"] = fund_ytd if as_of_year == year else np.nan
            row[f"{year} YTD Benchmark (%)"] = benchmark_ytd if as_of_year == year else np.nan
        rows.append(row)
    df = pd.DataFrame(rows)
    preferred = [
        "Fund Name", "Provider", "Share Class", "As Of Date", "1Y Return (%)",
        "3Y Return (%)", "5Y Return (%)", "3Y Cumulative (%)", "3Y Annualised (%)",
        "5Y Cumulative (%)", "5Y Annualised (%)", "Volatility (%)", "Mgmt Fee (%)",
        "1Y Benchmark (%)", "3Y Benchmark Cumulative (%)", "3Y Benchmark Annualised (%)",
        "5Y Benchmark Cumulative (%)", "5Y Benchmark Annualised (%)", "Benchmark Name",
        "Return Basis", "Source File", "Extraction Warnings",
    ]
    dynamic = []
    for year in years:
        dynamic.extend([f"{year} Return (%)", f"{year} Benchmark (%)"])
    for year in ytd_years:
        dynamic.extend([f"{year} YTD Return (%)", f"{year} YTD Benchmark (%)"])
    ordered = [col for col in preferred + dynamic if col in df.columns]
    ordered += [col for col in df.columns if col not in ordered]
    return df[ordered]


# Targeted overrides for the Streamlit FFS parser.
# These functions are appended after the baseline parser so that the tested
# provider-specific extraction rules replace the initial prototypes.


def _values_from_line(lines, pattern, start=0, stop=None, occurrence=1, count=20):
    regex = re.compile(pattern, re.I)
    stop = len(lines) if stop is None else min(stop, len(lines))
    seen = 0
    for i in range(start, stop):
        match = regex.search(lines[i])
        if not match:
            continue
        seen += 1
        if seen == occurrence:
            return _num_tokens(lines[i][match.end():])[:count]
    return []


def _numeric_block(lines, start, stop, limit=40):
    values = []
    for line in lines[start:stop]:
        parsed = _numeric_only_line(line)
        if parsed:
            values.extend(parsed)
        if len(values) >= limit:
            break
    return values[:limit]


def _calendar_from_header(lines, header_pattern, fund_pattern=None, benchmark_pattern=None, start=0):
    header_index = _marker_index(lines, header_pattern, start)
    if header_index is None:
        return {}
    years = [int(x) for x in re.findall(r"20\d{2}", lines[header_index])]
    if not years:
        return {}
    fund_values = []
    benchmark_values = []
    fund_pattern = fund_pattern or r"^Fund"
    benchmark_pattern = benchmark_pattern or r"^Benchmark"
    for i in range(header_index + 1, min(len(lines), header_index + 15)):
        if re.search(fund_pattern, lines[i], re.I):
            fund_values = _num_tokens(lines[i].split(None, 1)[1] if len(lines[i].split(None, 1)) > 1 else "")
            if not fund_values:
                fund_values = _row_values(lines, fund_pattern, start=i, count=len(years), lookahead=2)
        if re.search(benchmark_pattern, lines[i], re.I):
            benchmark_values = _num_tokens(lines[i].split(None, 1)[1] if len(lines[i].split(None, 1)) > 1 else "")
            if not benchmark_values:
                benchmark_values = _row_values(lines, benchmark_pattern, start=i, count=len(years), lookahead=2)
        if len(fund_values) >= len(years) and len(benchmark_values) >= len(years):
            break
    return {
        year: (
            fund_values[idx] if idx < len(fund_values) else np.nan,
            benchmark_values[idx] if idx < len(benchmark_values) else np.nan,
        )
        for idx, year in enumerate(years)
    }


def _set_period(record, period, fund_cum=np.nan, fund_ann=np.nan,
                bench_cum=np.nan, bench_ann=np.nan, basis=""):
    years = int(period.rstrip("Y"))
    raw_fund_cum = fund_cum
    raw_fund_ann = fund_ann
    raw_bench_cum = bench_cum
    raw_bench_ann = bench_ann
    if not np.isfinite(fund_cum) and np.isfinite(fund_ann):
        fund_cum = cumulative_from_annualized(fund_ann, years)
    if not np.isfinite(fund_ann) and np.isfinite(fund_cum):
        fund_ann = annualized_from_cumulative(fund_cum, years)
    if not np.isfinite(bench_cum) and np.isfinite(bench_ann):
        bench_cum = cumulative_from_annualized(bench_ann, years)
    if not np.isfinite(bench_ann) and np.isfinite(bench_cum):
        bench_ann = annualized_from_cumulative(bench_cum, years)

    record[f"{period} Cumulative (%)"] = fund_cum
    record[f"{period} Annualised (%)"] = fund_ann
    record[f"{period} Benchmark Cumulative (%)"] = bench_cum
    record[f"{period} Benchmark Annualised (%)"] = bench_ann
    # Legacy columns are deliberately annualised because the app's analysis
    # and existing charts use these names.
    record[f"{period} Return (%)"] = fund_ann
    record[f"{period} Benchmark (%)"] = bench_ann
    record[f"{period} Return Basis"] = basis
    record[f"{period} Return Source"] = (
        "FFS" if np.isfinite(raw_fund_cum) and np.isfinite(raw_fund_ann)
        else "FFS + calculated equivalent"
    )


def _parse_aham(lines, record):
    total_start = _marker_index(lines, r"Total Return \(%\)")
    annual_start = _marker_index(lines, r"Annualised Return \(%\)")
    calendar_start = _marker_index(lines, r"Calendar Year Return \(%\)")
    total_start = total_start if total_start is not None else 0
    annual_start = annual_start if annual_start is not None else 0

    total_fund = _values_from_line(lines, r"^Fund \(MYR\)", start=total_start, stop=annual_start or None, count=4)
    total_bench = _values_from_line(lines, r"^Benchmark \(MYR\)", start=total_start, stop=annual_start or None, count=4)
    annual_fund = _values_from_line(lines, r"^Fund \(MYR\)", start=annual_start, stop=calendar_start or None, count=4)
    annual_bench = _values_from_line(lines, r"^Benchmark \(MYR\)", start=annual_start, stop=calendar_start or None, count=4)

    if len(total_fund) >= 3:
        record["1Y Return (%)"] = total_fund[1]
        record["1Y Benchmark (%)"] = total_bench[1] if len(total_bench) >= 2 else np.nan
    if len(annual_fund) >= 3:
        _set_period(
            record, "3Y",
            fund_cum=total_fund[2] if len(total_fund) >= 3 else np.nan,
            fund_ann=annual_fund[1],
            bench_cum=total_bench[2] if len(total_bench) >= 3 else np.nan,
            bench_ann=annual_bench[1] if len(annual_bench) >= 2 else np.nan,
            basis="FFS reports cumulative 3Y and annualised 3Y",
        )
        _set_period(
            record, "5Y",
            fund_ann=annual_fund[2],
            bench_ann=annual_bench[2] if len(annual_bench) >= 3 else np.nan,
            basis="FFS reports annualised 5Y; cumulative equivalent calculated",
        )
    else:
        _append_warning(record, "AHAM annualised return table was not detected")

    if calendar_start is not None:
        calendar = _calendar_from_header(
            lines, r"Calendar Year Return \(%\)",
            fund_pattern=r"^Fund \(MYR\)",
            benchmark_pattern=r"^Benchmark \(MYR\)",
            start=calendar_start,
        )
        fy = _values_from_line(lines, r"^Fund \(MYR\)", start=calendar_start, count=4)
        by = _values_from_line(lines, r"^Benchmark \(MYR\)", start=calendar_start, count=4)
        years = [int(x) for x in re.findall(r"20\d{2}", lines[calendar_start])]
        record["Calendar"] = {
            year: (fy[idx] if idx < len(fy) else np.nan, by[idx] if idx < len(by) else np.nan)
            for idx, year in enumerate(years)
        }
        record["YTD"] = (fy[0] if fy else np.nan, by[0] if by else np.nan)
    return record


def _parse_amfunds(lines, record):
    cumulative_start = _marker_index(lines, r"Cumulative Return \(%\)")
    annual_start = _marker_index(lines, r"Annualised Return \(%\)")
    calendar_start = _marker_index(lines, r"Calendar Year Return \(%\)")
    cumulative_start = cumulative_start if cumulative_start is not None else 0
    annual_start = annual_start if annual_start is not None else len(lines)

    fund_cum = _values_from_line(lines, r"^Fund\s+", start=cumulative_start, stop=annual_start, count=6)
    bench_cum = _values_from_line(lines, r"^\*Benchmark", start=cumulative_start, stop=annual_start, count=6)
    fund_ann = _values_from_line(lines, r"^Fund\s+", start=annual_start, stop=calendar_start or None, count=4)
    bench_ann = _values_from_line(lines, r"^\*Benchmark", start=annual_start, stop=calendar_start or None, count=4)

    if len(fund_cum) >= 6:
        record["1Y Return (%)"] = fund_cum[3]
        record["1Y Benchmark (%)"] = bench_cum[3] if len(bench_cum) >= 4 else np.nan
    if len(fund_ann) >= 2:
        # AmBond labels its annualised rows as 3Y, 5Y, 10Y, Since Inception.
        _set_period(
            record, "3Y",
            fund_cum=fund_cum[4] if len(fund_cum) >= 5 else np.nan,
            fund_ann=fund_ann[0],
            bench_cum=bench_cum[4] if len(bench_cum) >= 5 else np.nan,
            bench_ann=bench_ann[0] if len(bench_ann) >= 1 else np.nan,
            basis="FFS reports cumulative and annualised 3Y",
        )
        _set_period(
            record, "5Y",
            fund_cum=fund_cum[5] if len(fund_cum) >= 6 else np.nan,
            fund_ann=fund_ann[1],
            bench_cum=bench_cum[5] if len(bench_cum) >= 6 else np.nan,
            bench_ann=bench_ann[1] if len(bench_ann) >= 2 else np.nan,
            basis="FFS reports cumulative and annualised 5Y",
        )
    else:
        _append_warning(record, "AmBond cumulative or annualised table was not detected")

    if calendar_start is not None:
        years = [int(x) for x in re.findall(r"20\d{2}", lines[calendar_start])]
        fy = _values_from_line(lines, r"^Fund\s+", start=calendar_start, count=len(years))
        by = _values_from_line(lines, r"^\*Benchmark", start=calendar_start, count=len(years))
        record["Calendar"] = {
            year: (fy[idx] if idx < len(fy) else np.nan, by[idx] if idx < len(by) else np.nan)
            for idx, year in enumerate(years)
        }
        record["YTD"] = (fy[0] if fy else np.nan, by[0] if by else np.nan)
    return record


def _parse_eastspring(lines, record):
    performance_start = _marker_index(lines, r"^PERFORMANCE TABLE$")
    if performance_start is None:
        _append_warning(record, "Eastspring performance table was not detected")
        return record
    fund_marker = _marker_index(lines, r"^Fund$", performance_start)
    benchmark_marker = _marker_index(lines, r"^Benchmark", performance_start)
    start = max(fund_marker or performance_start, benchmark_marker or performance_start) + 1
    stop = _marker_index(lines, r"10 years", start) or min(len(lines), start + 40)
    values = _numeric_block(lines, start, stop, limit=20)
    pairs = [(values[i], values[i + 1]) for i in range(0, len(values) - 1, 2)]
    if len(pairs) >= 5:
        record["1Y Return (%)"] = pairs[2][0]
        record["1Y Benchmark (%)"] = pairs[2][1]
        _set_period(record, "3Y", fund_cum=pairs[3][0], bench_cum=pairs[3][1], basis="FFS reports cumulative 3Y; annualised equivalent calculated")
        _set_period(record, "5Y", fund_cum=pairs[4][0], bench_cum=pairs[4][1], basis="FFS reports cumulative 5Y; annualised equivalent calculated")
        if len(pairs) >= 7:
            record["YTD"] = pairs[6]
    else:
        _append_warning(record, "Eastspring numeric performance rows were not detected")

    header = _marker_index(lines, r"^Year\.")
    if header is not None:
        years = [int(x) for x in re.findall(r"20\d{2}", lines[header])]
        fund_line = _marker_index(lines, r"^Annual Fund Performance", header)
        bench_label = _marker_index(lines, r"^Annual Benchmark", header)
        bench_line = (bench_label + 1) if bench_label is not None and bench_label + 1 < len(lines) else None
        fy = _num_tokens(lines[fund_line].split(".", 1)[-1]) if fund_line is not None else []
        by = _num_tokens(lines[bench_line]) if bench_line is not None else []
        record["Calendar"] = {
            year: (fy[idx] if idx < len(fy) else np.nan, by[idx] if idx < len(by) else np.nan)
            for idx, year in enumerate(years)
        }
    return record


def _parse_manulife(lines, record):
    start = _marker_index(lines, r"Total return over the following periods") or 0
    fund = _values_from_line(lines, r"^Fund A \(RM Hedged\) Class", start=start, count=7)
    bench = _values_from_line(lines, r"^Benchmark in USD", start=start, count=7)
    if len(fund) >= 6:
        record["1Y Return (%)"] = fund[3]
        record["1Y Benchmark (%)"] = bench[3] if len(bench) >= 4 else np.nan
        _set_period(record, "3Y", fund_cum=fund[4], bench_cum=bench[4] if len(bench) >= 5 else np.nan, basis="FFS reports cumulative 3Y; annualised equivalent calculated")
        _set_period(record, "5Y", fund_cum=fund[5], bench_cum=bench[5] if len(bench) >= 6 else np.nan, basis="FFS reports cumulative 5Y; annualised equivalent calculated")
        record["YTD"] = (fund[2], bench[2] if len(bench) >= 3 else np.nan)
    else:
        _append_warning(record, "Manulife total-return table was not detected")
    calendar_start = _marker_index(lines, r"Calendar year returns")
    if calendar_start is not None:
        years = [int(x) for x in re.findall(r"20\d{2}", lines[calendar_start])]
        fy = _values_from_line(lines, r"^Fund A \(RM Hedged\) Class", start=calendar_start, count=len(years))
        by = _values_from_line(lines, r"^Benchmark in USD", start=calendar_start, count=len(years))
        record["Calendar"] = {
            year: (fy[idx] if idx < len(fy) else np.nan, by[idx] if idx < len(by) else np.nan)
            for idx, year in enumerate(years)
        }
    return record


def _parse_principal(lines, record):
    start = _marker_index(lines, r"Cumulative Performance \(%\)") or 0
    fund = _values_from_line(lines, r"^Fund\s+", start=start, count=8)
    bench = _values_from_line(lines, r"^Benchmark\s+", start=start, count=8)
    if len(fund) >= 7:
        record["1Y Return (%)"] = fund[4]
        record["1Y Benchmark (%)"] = bench[4] if len(bench) >= 5 else np.nan
        _set_period(record, "3Y", fund_cum=fund[5], bench_cum=bench[5] if len(bench) >= 6 else np.nan, basis="FFS reports cumulative 3Y; annualised equivalent calculated")
        _set_period(record, "5Y", fund_cum=fund[6], bench_cum=bench[6] if len(bench) >= 7 else np.nan, basis="FFS reports cumulative 5Y; annualised equivalent calculated")
    else:
        _append_warning(record, "Principal cumulative-performance table was not detected")
    calendar_start = _marker_index(lines, r"Calendar Year Returns")
    if calendar_start is not None:
        years = [int(x) for x in re.findall(r"20\d{2}", lines[calendar_start])]
        fy = _values_from_line(lines, r"^Fund\s+", start=calendar_start, count=len(years))
        by = _values_from_line(lines, r"^Benchmark\s+", start=calendar_start, count=len(years))
        record["Calendar"] = {
            year: (fy[idx] if idx < len(fy) else np.nan, by[idx] if idx < len(by) else np.nan)
            for idx, year in enumerate(years)
        }
    return record


def _parse_maybank(lines, record):
    # Maybank's standard extraction keeps the entire performance rows on one line.
    total_start = _marker_index(lines, r"^Total Return") or 0
    annual_start = _marker_index(lines, r"^Annual Management Fee") or 0
    total = _values_from_line(lines, r"^MGSEIF \(MYRH\)", start=total_start, stop=annual_start, occurrence=1, count=8)
    bench_total = _values_from_line(lines, r"^Benchmark", start=total_start, stop=annual_start, occurrence=1, count=8)
    annual = _values_from_line(lines, r"^MGSEIF \(MYRH\)", start=annual_start, count=8)
    bench_annual = _values_from_line(lines, r"^Benchmark", start=annual_start, count=8)
    if len(total) >= 7:
        record["1Y Return (%)"] = total[4]
        record["1Y Benchmark (%)"] = bench_total[4] if len(bench_total) >= 5 else np.nan
        _set_period(record, "3Y", fund_cum=total[5], fund_ann=annual[0] if len(annual) >= 1 else np.nan, bench_cum=bench_total[5] if len(bench_total) >= 6 else np.nan, bench_ann=bench_annual[0] if len(bench_annual) >= 1 else np.nan, basis="FFS reports cumulative and annualised 3Y")
        _set_period(record, "5Y", fund_cum=total[6], fund_ann=annual[1] if len(annual) >= 2 else np.nan, bench_cum=bench_total[6] if len(bench_total) >= 7 else np.nan, bench_ann=bench_annual[1] if len(bench_annual) >= 2 else np.nan, basis="FFS reports cumulative and annualised 5Y")
        record["YTD"] = (total[0], bench_total[0] if bench_total else np.nan)
    else:
        _append_warning(record, "Maybank MYRH total-return row was not detected")
    calendar_start = _marker_index(lines, r"Calendar Year Return")
    if calendar_start is not None:
        years = [int(x) for x in re.findall(r"20\d{2}", lines[annual_start])]
        # Annual rows are headed 3Y, 5Y, SI, 2025, 2024, 2023, 2022, 2021.
        if len(annual) >= 3:
            record["Calendar"] = {
                year: (annual[idx + 3] if idx + 3 < len(annual) else np.nan, bench_annual[idx + 3] if idx + 3 < len(bench_annual) else np.nan)
                for idx, year in enumerate(years)
            }
    return record


def _parse_ta(lines, record):
    start = _marker_index(lines, r"Cumulative Fund Performance") or 0
    fund = _values_from_line(lines, r"^TA Global T echnology Fund MYR Hdg", start=start, count=8)
    bench = _values_from_line(lines, r"^TAGTF Benchmark", start=start, count=8)
    if len(fund) >= 6:
        record["1Y Return (%)"] = fund[2]
        record["1Y Benchmark (%)"] = bench[2] if len(bench) >= 3 else np.nan
        _set_period(record, "3Y", fund_cum=fund[3], bench_cum=bench[3] if len(bench) >= 4 else np.nan, basis="FFS reports cumulative 3Y; annualised equivalent calculated")
        _set_period(record, "5Y", fund_cum=fund[4], bench_cum=bench[4] if len(bench) >= 5 else np.nan, basis="FFS reports cumulative 5Y; annualised equivalent calculated")
        record["YTD"] = (fund[5], bench[5] if len(bench) >= 6 else np.nan)
    else:
        _append_warning(record, "TA MYR Hedged cumulative-performance row was not detected")
    annual_start = _marker_index(lines, r"Annual Fund Performance \(%\)")
    if annual_start is not None:
        years = [int(x) for x in re.findall(r"20\d{2}", lines[annual_start])]
        af = _values_from_line(lines, r"^TA Global T echnology Fund MYR Hdg", start=annual_start, count=len(years))
        ab = _values_from_line(lines, r"^TAGTF Benchmark", start=annual_start, count=len(years))
        record["Calendar"] = {
            year: (af[idx] if idx < len(af) else np.nan, ab[idx] if idx < len(ab) else np.nan)
            for idx, year in enumerate(years)
        }
    _append_warning(record, "TA volatility value is not exposed in the text layer; verify or enter it manually")
    return record


def parse_ffs_bytes(pdf_bytes, filename):
    # Use layout mode for the AHAM and AmBond tables, whose labels and values
    # are positioned in two-column fact sheets. Use standard mode for the
    # remaining templates, especially TA, where layout mode loses the table.
    standard_text = ""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes), strict=False)
        standard_text = "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:
        standard_text = ""
    provider_hint = _identify_provider(filename, standard_text)
    text = standard_text
    if provider_hint in {"AHAM", "AmFunds"}:
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes), strict=False)
            text = "\n".join((page.extract_text(extraction_mode="layout") or "") for page in reader.pages)
        except Exception:
            text = standard_text
    if len(re.sub(r"\s+", "", text)) < 80:
        temp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as handle:
                handle.write(pdf_bytes)
                temp_path = handle.name
            text = extract_text(temp_path)
        finally:
            if temp_path and os.path.exists(temp_path):
                os.unlink(temp_path)
    if not text or len(re.sub(r"\s+", "", text)) < 80:
        raise ValueError("No usable selectable text was found. This appears to be an image-only PDF and cannot be processed under the no-OCR policy.")
    return parse_ffs_text(text, filename)


def records_to_dataframe(records):
    rows = []
    years = sorted({year for record in records for year in record.get("Calendar", {})})
    ytd_years = sorted({int(str(record.get("As Of Date", ""))[-4:]) for record in records if record.get("YTD") and any(np.isfinite(x) for x in record.get("YTD", (np.nan, np.nan)))})
    for record in records:
        row = {key: value for key, value in record.items() if key not in {"Calendar", "YTD"}}
        for year in years:
            fund, benchmark = record.get("Calendar", {}).get(year, (np.nan, np.nan))
            row[f"{year} Return (%)"] = fund
            row[f"{year} Benchmark (%)"] = benchmark
        as_of_year_match = re.search(r"(20\d{2})", str(record.get("As Of Date", "")))
        as_of_year = int(as_of_year_match.group(1)) if as_of_year_match else None
        fund_ytd, benchmark_ytd = record.get("YTD", (np.nan, np.nan))
        for year in ytd_years:
            row[f"{year} YTD Return (%)"] = fund_ytd if as_of_year == year else np.nan
            row[f"{year} YTD Benchmark (%)"] = benchmark_ytd if as_of_year == year else np.nan
        rows.append(row)
    df = pd.DataFrame(rows)
    preferred = [
        "Fund Name", "Provider", "Share Class", "As Of Date", "1Y Return (%)",
        "3Y Return (%)", "3Y Cumulative (%)", "3Y Annualised (%)", "3Y Return Source",
        "5Y Return (%)", "5Y Cumulative (%)", "5Y Annualised (%)", "5Y Return Source",
        "Volatility (%)", "Mgmt Fee (%)", "1Y Benchmark (%)",
        "3Y Benchmark (%)", "3Y Benchmark Cumulative (%)", "3Y Benchmark Annualised (%)",
        "5Y Benchmark (%)", "5Y Benchmark Cumulative (%)", "5Y Benchmark Annualised (%)",
        "Benchmark Name", "Return Basis", "Source File", "Extraction Warnings",
    ]
    dynamic = []
    for year in years:
        dynamic.extend([f"{year} Return (%)", f"{year} Benchmark (%)"])
    for year in ytd_years:
        dynamic.extend([f"{year} YTD Return (%)", f"{year} YTD Benchmark (%)"])
    ordered = [col for col in preferred + dynamic if col in df.columns]
    ordered += [col for col in df.columns if col not in ordered]
    return df[ordered]


# Alias used by the app when it needs to identify whether a return is printed
# or calculated after a manual edit.
def analysis_return_series(df):
    result = df["1Y Return (%)"].copy() if "1Y Return (%)" in df.columns else pd.Series(np.nan, index=df.index)
    if "3Y Annualised (%)" in df.columns:
        result = result.fillna(df["3Y Annualised (%)"])
    if "5Y Annualised (%)" in df.columns:
        result = result.fillna(df["5Y Annualised (%)"])
    if "3Y Return (%)" in df.columns:
        result = result.fillna(df["3Y Return (%)"])
    if "5Y Return (%)" in df.columns:
        result = result.fillna(df["5Y Return (%)"])
    return result


def _metadata(text, provider):
    compact = re.sub(r"\s+", " ", text)
    volatility = _first_float([
        r"Volatility Factor.*?is\s+(\d+(?:\.\d+)?)",
        r"volatility.{0,100}?(\d+(?:\.\d+)?)",
    ], compact)
    if provider == "TA":
        volatility = np.nan
    fee = _first_float([
        r"Annual Management Fee.{0,100}?(\d+(?:\.\d+)?)\s*%",
        r"Management Fee.{0,100}?(\d+(?:\.\d+)?)\s*%",
    ], compact)
    if provider == "AHAM" and not np.isfinite(fee):
        fee = _first_float([r"Max\s+(\d+(?:\.\d+)?)\s*%\s*per annum"], compact)
    date = ""
    date_match = re.search(r"as (?:at|of)\s+(\d{1,2}[ /-][A-Za-z0-9]+[ /-]\d{2,4})", compact, re.I)
    if date_match:
        date = date_match.group(1)
    else:
        date_match = re.search(r"As of Date:\s*(\d{1,2}/\d{1,2}/\d{4})", compact, re.I)
        if date_match:
            date = date_match.group(1)
        else:
            date_match = re.search(r"as at\s+(\d{1,2}[A-Za-z -]+\d{4})", compact, re.I)
            if date_match:
                date = date_match.group(1)
    benchmark_patterns = [
        r"MSCI AC Asia Pacific ex Japan High Dividend Yield Index",
        r"BPAM Corporates All Bond Index",
        r"MSCI Emerging Markets Net Total Return Index",
        r"Dow Jones Islamic Market World Index",
        r"MSCI World/Health Care NR USD Index",
        r"42% S&P 500 \+ 36% MSCI Europe \+ 12% MSCI Japan \+ 10% CIMB Bank 1-Month Fixed Deposit Rate",
        r"MSCI ACWI Information Technology Index \+ MSCI ACWI Communication Services Index",
        r"Dow Jones Islamic Market World Index",
    ]
    benchmark = ""
    for pattern in benchmark_patterns:
        match = re.search(pattern, compact, re.I)
        if match:
            benchmark = match.group(0)
            break
    return volatility, fee, date, benchmark


# Ensure the DataFrame fields required by the existing app are always present.
_original_base_record = _base_record

def _base_record(filename, text, provider, fund_name, share_class):
    record = _original_base_record(filename, text, provider, fund_name, share_class)
    record.update({
        "3Y Return Source": "",
        "5Y Return Source": "",
        "3Y Benchmark (%)": np.nan,
        "5Y Benchmark (%)": np.nan,
    })
    return record

_original_parse_ffs_text = parse_ffs_text

def parse_ffs_text(text, filename):
    text = _clean_text(text)
    lines = _lines(text)
    provider = _identify_provider(filename, text)
    fund_name, share_class = _identity(filename, text, provider)
    record = _base_record(filename, text, provider, fund_name, share_class)
    parser = {
        "AHAM": _parse_aham,
        "AmFunds": _parse_amfunds,
        "Eastspring": _parse_eastspring,
        "Manulife": _parse_manulife,
        "Maybank": _parse_maybank,
        "Principal": _parse_principal,
        "TA": _parse_ta,
    }.get(provider)
    if parser:
        record = parser(lines, record)
    else:
        _append_warning(record, "Unknown FFS template; manual extraction review required")
    return record


# Keep benchmark legacy columns synchronized with the annualised values used by
# the existing analysis/report/chart code.
_original_set_period = _set_period

def _set_period(record, period, fund_cum=np.nan, fund_ann=np.nan,
                bench_cum=np.nan, bench_ann=np.nan, basis=""):
    _original_set_period(record, period, fund_cum, fund_ann, bench_cum, bench_ann, basis)
    record[f"{period} Benchmark (%)"] = record.get(f"{period} Benchmark Annualised (%)", np.nan)
    record[f"{period} Return (%)"] = record.get(f"{period} Annualised (%)", np.nan)


def parse_ffs_bytes(pdf_bytes, filename):
    standard_text = ""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes), strict=False)
        standard_text = "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:
        standard_text = ""
    provider_hint = _identify_provider(filename, standard_text)
    text = standard_text
    if provider_hint in {"AHAM", "AmFunds"}:
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes), strict=False)
            text = "\n".join((page.extract_text(extraction_mode="layout") or "") for page in reader.pages)
        except Exception:
            text = standard_text
    if len(re.sub(r"\s+", "", text)) < 80:
        temp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as handle:
                handle.write(pdf_bytes)
                temp_path = handle.name
            text = extract_text(temp_path)
        finally:
            if temp_path and os.path.exists(temp_path):
                os.unlink(temp_path)
    if not text or len(re.sub(r"\s+", "", text)) < 80:
        raise ValueError("No usable selectable text was found. This appears to be an image-only PDF and cannot be processed under the no-OCR policy.")
    return parse_ffs_text(text, filename)


def records_to_dataframe(records):
    rows = []
    years = sorted({year for record in records for year in record.get("Calendar", {})})
    ytd_years = sorted({int(str(record.get("As Of Date", ""))[-4:]) for record in records if record.get("YTD") and any(np.isfinite(x) for x in record.get("YTD", (np.nan, np.nan)))})
    for record in records:
        row = {key: value for key, value in record.items() if key not in {"Calendar", "YTD"}}
        for year in years:
            fund, benchmark = record.get("Calendar", {}).get(year, (np.nan, np.nan))
            row[f"{year} Return (%)"] = fund
            row[f"{year} Benchmark (%)"] = benchmark
        as_of_year_match = re.search(r"(20\d{2})", str(record.get("As Of Date", "")))
        as_of_year = int(as_of_year_match.group(1)) if as_of_year_match else None
        fund_ytd, benchmark_ytd = record.get("YTD", (np.nan, np.nan))
        for year in ytd_years:
            row[f"{year} YTD Return (%)"] = fund_ytd if as_of_year == year else np.nan
            row[f"{year} YTD Benchmark (%)"] = benchmark_ytd if as_of_year == year else np.nan
        rows.append(row)
    df = pd.DataFrame(rows)
    preferred = [
        "Fund Name", "Provider", "Share Class", "As Of Date", "1Y Return (%)",
        "3Y Return (%)", "3Y Cumulative (%)", "3Y Annualised (%)", "3Y Return Source",
        "5Y Return (%)", "5Y Cumulative (%)", "5Y Annualised (%)", "5Y Return Source",
        "Volatility (%)", "Mgmt Fee (%)", "1Y Benchmark (%)", "3Y Benchmark (%)",
        "3Y Benchmark Cumulative (%)", "3Y Benchmark Annualised (%)", "5Y Benchmark (%)",
        "5Y Benchmark Cumulative (%)", "5Y Benchmark Annualised (%)", "Benchmark Name",
        "Return Basis", "Source File", "Extraction Warnings",
    ]
    dynamic = []
    for year in years:
        dynamic.extend([f"{year} Return (%)", f"{year} Benchmark (%)"])
    for year in ytd_years:
        dynamic.extend([f"{year} YTD Return (%)", f"{year} YTD Benchmark (%)"])
    ordered = [col for col in preferred + dynamic if col in df.columns]
    ordered += [col for col in df.columns if col not in ordered]
    return df[ordered]


def _analysis_return_series(df):
    return analysis_return_series(df)


def _normalize_analysis_columns(df):
    df = df.copy()
    for col in ["1Y Return (%)", "3Y Return (%)", "5Y Return (%)", "3Y Annualised (%)", "5Y Annualised (%)", "Volatility (%)", "Mgmt Fee (%)", "1Y Benchmark (%)", "3Y Benchmark (%)", "5Y Benchmark (%)"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    df["1Y Return (%)"] = _analysis_return_series(df)
    if "3Y Annualised (%)" in df.columns:
        df["3Y Return (%)"] = pd.to_numeric(df["3Y Annualised (%)"], errors="coerce")
    if "5Y Annualised (%)" in df.columns:
        df["5Y Return (%)"] = pd.to_numeric(df["5Y Annualised (%)"], errors="coerce")
    return df


# Final corrections based on the actual pypdf output from the eight uploaded
# samples. Several fact sheets concatenate the left-hand metadata column with
# the performance table, so labels must not be anchored at column zero.


def _values_after_label(lines, pattern, start=0, stop=None, occurrence=1, count=20, lookahead=8):
    regex = re.compile(pattern, re.I)
    stop = len(lines) if stop is None else min(stop, len(lines))
    seen = 0
    for i in range(start, stop):
        match = regex.search(lines[i])
        if not match:
            continue
        seen += 1
        if seen != occurrence:
            continue
        values = _num_tokens(lines[i][match.end():])
        if len(values) >= count:
            return values[:count]
        for j in range(i + 1, min(stop, i + lookahead + 1)):
            candidate = lines[j].strip()
            # Fact sheets often put the numeric row on a following '(%)' line.
            if candidate.startswith("(%)") or candidate.startswith("%"):
                values = _num_tokens(candidate)
                if len(values) >= count:
                    return values[:count]
            numeric = _numeric_only_line(candidate)
            if numeric:
                values.extend(numeric)
                if len(values) >= count:
                    return values[:count]
        return values[:count]
    return []


def _parse_aham(lines, record):
    total_start = _marker_index(lines, r"Total Return \(%\)")
    annual_start = _marker_index(lines, r"Annualised Return \(%\)")
    calendar_start = _marker_index(lines, r"Calendar Year Return \(%\)")
    total_start = total_start if total_start is not None else 0
    annual_start = annual_start if annual_start is not None else len(lines)
    fund_total = _values_after_label(lines, r"Fund \(MYR\)", start=total_start, stop=annual_start, count=4)
    bench_total = _values_after_label(lines, r"Benchmark \(MYR\)", start=total_start, stop=annual_start, count=4)
    fund_annual = _values_after_label(lines, r"Fund \(MYR\)", start=annual_start, stop=calendar_start, count=4)
    bench_annual = _values_after_label(lines, r"Benchmark \(MYR\)", start=annual_start, stop=calendar_start, count=4)

    if len(fund_total) >= 3:
        record["1Y Return (%)"] = fund_total[1]
        record["1Y Benchmark (%)"] = bench_total[1] if len(bench_total) >= 2 else np.nan
    if len(fund_annual) >= 3:
        _set_period(record, "3Y", fund_cum=fund_total[2] if len(fund_total) >= 3 else np.nan, fund_ann=fund_annual[1], bench_cum=bench_total[2] if len(bench_total) >= 3 else np.nan, bench_ann=bench_annual[1] if len(bench_annual) >= 2 else np.nan, basis="FFS reports cumulative and annualised 3Y")
        _set_period(record, "5Y", fund_ann=fund_annual[2], bench_ann=bench_annual[2] if len(bench_annual) >= 3 else np.nan, basis="FFS reports annualised 5Y; cumulative equivalent calculated")
    else:
        _append_warning(record, "AHAM annualised return table was not detected")

    if calendar_start is not None:
        years = [int(x) for x in re.findall(r"20\d{2}", lines[calendar_start])]
        fund_values = _values_after_label(lines, r"Fund \(MYR\)", start=calendar_start, count=len(years))
        bench_values = _values_after_label(lines, r"Benchmark \(MYR\)", start=calendar_start, count=len(years))
        record["Calendar"] = {
            year: (fund_values[idx] if idx < len(fund_values) else np.nan, bench_values[idx] if idx < len(bench_values) else np.nan)
            for idx, year in enumerate(years)
        }
        record["YTD"] = (fund_values[0] if fund_values else np.nan, bench_values[0] if bench_values else np.nan)
    return record


def _parse_eastspring(lines, record):
    performance_start = _marker_index(lines, r"^PERFORMANCE TABLE$")
    if performance_start is None:
        _append_warning(record, "Eastspring performance table was not detected")
        return record
    fund_marker = _marker_index(lines, r"^Fund$", performance_start)
    benchmark_marker = _marker_index(lines, r"^Benchmark", performance_start)
    start = max(fund_marker or performance_start, benchmark_marker or performance_start) + 1
    stop = _marker_index(lines, r"10 years", start) or min(len(lines), start + 40)
    values = _numeric_block(lines, start, stop, limit=20)
    pairs = [(values[i], values[i + 1]) for i in range(0, len(values) - 1, 2)]
    if len(pairs) >= 5:
        record["1Y Return (%)"] = pairs[2][0]
        record["1Y Benchmark (%)"] = pairs[2][1]
        _set_period(record, "3Y", fund_cum=pairs[3][0], bench_cum=pairs[3][1], basis="FFS reports cumulative 3Y; annualised equivalent calculated")
        _set_period(record, "5Y", fund_cum=pairs[4][0], bench_cum=pairs[4][1], basis="FFS reports cumulative 5Y; annualised equivalent calculated")
        record["YTD"] = pairs[6] if len(pairs) > 6 else (np.nan, np.nan)
    else:
        _append_warning(record, "Eastspring numeric performance rows were not detected")

    header = _marker_index(lines, r"^Year\.")
    if header is not None:
        years = [int(x) for x in re.findall(r"20\d{2}", lines[header])]
        fund_line = _marker_index(lines, r"^Annual Fund Performance", header)
        bench_label = _marker_index(lines, r"^Annual Benchmark", header)
        bench_line = None
        if bench_label is not None:
            for j in range(bench_label, min(len(lines), bench_label + 4)):
                if _num_tokens(lines[j]):
                    bench_line = j
                    break
        fy = _num_tokens(lines[fund_line].split(".", 1)[-1]) if fund_line is not None else []
        by = _num_tokens(lines[bench_line]) if bench_line is not None else []
        record["Calendar"] = {
            year: (fy[idx] if idx < len(fy) else np.nan, by[idx] if idx < len(by) else np.nan)
            for idx, year in enumerate(years)
        }
    return record


def _parse_maybank(lines, record):
    total_start = _marker_index(lines, r"Total Return") or 0
    annual_header = _marker_index(lines, r"^Annual Management Fee") or len(lines)
    fund_total = _values_after_label(lines, r"^MGSEIF \(MYRH\)", start=total_start, stop=annual_header, count=8)
    bench_total = _values_after_label(lines, r"^Benchmark", start=total_start, stop=annual_header, count=8)
    fund_annual = _values_after_label(lines, r"^MGSEIF \(MYRH\)", start=annual_header, count=8)
    bench_annual = _values_after_label(lines, r"^Benchmark", start=annual_header, count=8)
    if len(fund_total) >= 7:
        record["1Y Return (%)"] = fund_total[4]
        record["1Y Benchmark (%)"] = bench_total[4] if len(bench_total) >= 5 else np.nan
        _set_period(record, "3Y", fund_cum=fund_total[5], fund_ann=fund_annual[0] if fund_annual else np.nan, bench_cum=bench_total[5] if len(bench_total) >= 6 else np.nan, bench_ann=bench_annual[0] if bench_annual else np.nan, basis="FFS reports cumulative and annualised 3Y")
        _set_period(record, "5Y", fund_cum=fund_total[6], fund_ann=fund_annual[1] if len(fund_annual) >= 2 else np.nan, bench_cum=bench_total[6] if len(bench_total) >= 7 else np.nan, bench_ann=bench_annual[1] if len(bench_annual) >= 2 else np.nan, basis="FFS reports cumulative and annualised 5Y")
        record["YTD"] = (fund_total[0], bench_total[0] if bench_total else np.nan)
    else:
        _append_warning(record, "Maybank MYRH total-return row was not detected")
    annual_line = _marker_index(lines, r"^Annual Management Fee")
    if annual_line is not None:
        years = [int(x) for x in re.findall(r"20\d{2}", lines[annual_line])]
        record["Calendar"] = {
            year: (fund_annual[idx + 3] if idx + 3 < len(fund_annual) else np.nan, bench_annual[idx + 3] if idx + 3 < len(bench_annual) else np.nan)
            for idx, year in enumerate(years)
        }
    return record


def _parse_manulife(lines, record):
    start = _marker_index(lines, r"Total return over the following periods") or 0
    fund = _values_after_label(lines, r"^Fund A \(RM Hedged\) Class", start=start, count=7)
    bench = _values_after_label(lines, r"^Benchmark in USD", start=start, count=7)
    if len(fund) >= 6:
        record["1Y Return (%)"] = fund[3]
        record["1Y Benchmark (%)"] = bench[3] if len(bench) >= 4 else np.nan
        _set_period(record, "3Y", fund_cum=fund[4], bench_cum=bench[4] if len(bench) >= 5 else np.nan, basis="FFS reports cumulative 3Y; annualised equivalent calculated")
        _set_period(record, "5Y", fund_cum=fund[5], bench_cum=bench[5] if len(bench) >= 6 else np.nan, basis="FFS reports cumulative 5Y; annualised equivalent calculated")
        record["YTD"] = (fund[2], bench[2] if len(bench) >= 3 else np.nan)
    else:
        _append_warning(record, "Manulife total-return table was not detected")
    calendar_start = _marker_index(lines, r"Calendar year returns")
    if calendar_start is not None:
        years = [int(x) for x in re.findall(r"20\d{2}", lines[calendar_start + 1])]
        fy = _values_after_label(lines, r"^Fund A \(RM Hedged\) Class", start=calendar_start, count=len(years))
        by = _values_after_label(lines, r"^Benchmark in USD", start=calendar_start, count=len(years))
        record["Calendar"] = {
            year: (fy[idx] if idx < len(fy) else np.nan, by[idx] if idx < len(by) else np.nan)
            for idx, year in enumerate(years)
        }
    return record


def _parse_ta(lines, record):
    cumulative_start = _marker_index(lines, r"Cumulative Fund Performance") or 0
    fund = _values_after_label(lines, r"^TA Global T echnology Fund MYR Hdg", start=cumulative_start, count=8)
    benchmark = _values_after_label(lines, r"^TAGTF Benchmark", start=cumulative_start, count=8)
    if len(fund) >= 6:
        record["1Y Return (%)"] = fund[2]
        record["1Y Benchmark (%)"] = benchmark[2] if len(benchmark) >= 3 else np.nan
        _set_period(record, "3Y", fund_cum=fund[3], bench_cum=benchmark[3] if len(benchmark) >= 4 else np.nan, basis="FFS reports cumulative 3Y; annualised equivalent calculated")
        _set_period(record, "5Y", fund_cum=fund[4], bench_cum=benchmark[4] if len(benchmark) >= 5 else np.nan, basis="FFS reports cumulative 5Y; annualised equivalent calculated")
        record["YTD"] = (fund[5], benchmark[5] if len(benchmark) >= 6 else np.nan)
    else:
        _append_warning(record, "TA MYR Hedged cumulative-performance row was not detected")
    annual_start = _marker_index(lines, r"^Annual Fund Performance")
    if annual_start is not None:
        header_line = annual_start + 1
        years = [int(x) for x in re.findall(r"20\d{2}", lines[header_line])] if header_line < len(lines) else [2025, 2024, 2023, 2022, 2021]
        annual_fund = _values_after_label(lines, r"^TA Global T echnology Fund MYR Hdg", start=annual_start, count=len(years))
        annual_bench = _values_after_label(lines, r"^TAGTF Benchmark", start=annual_start, count=len(years))
        record["Calendar"] = {
            year: (annual_fund[idx] if idx < len(annual_fund) else np.nan, annual_bench[idx] if idx < len(annual_bench) else np.nan)
            for idx, year in enumerate(years)
        }
    _append_warning(record, "TA volatility badge is not exposed cleanly in the text layer; verify manually")
    return record


def _metadata(text, provider):
    compact = re.sub(r"\s+", " ", text)
    volatility = np.nan
    volatility_patterns = [
        r"Volatility Factor.*?for this Fund is\s+(\d+(?:\.\d+)?)",
        r"Volatility Factor.*?is\s+(\d+(?:\.\d+)?)",
    ]
    if provider != "TA":
        volatility = _first_float(volatility_patterns, compact)
    fee = _first_float([
        r"Annual Management Fee.{0,100}?(\d+(?:\.\d+)?)\s*%",
        r"Management Fee.{0,100}?(\d+(?:\.\d+)?)\s*%",
    ], compact)
    if provider == "AHAM" and not np.isfinite(fee):
        fee = _first_float([r"Max\s+(\d+(?:\.\d+)?)\s*%\s*per annum"], compact)
    date = ""
    for pattern in [
        r"as (?:at|of)\s+(\d{1,2}[ /-][A-Za-z0-9]+[ /-]\d{2,4})",
        r"As of Date:\s*(\d{1,2}/\d{1,2}/\d{4})",
    ]:
        match = re.search(pattern, compact, re.I)
        if match:
            date = match.group(1)
            break
    benchmark = ""
    for pattern in [
        r"MSCI AC Asia Pacific ex Japan High Dividend Yield Index",
        r"BPAM Corporates All Bond Index",
        r"MSCI Emerging Markets Net Total Return Index",
        r"Dow Jones Islamic Market World Index",
        r"MSCI World / Health Care NR USD Index",
        r"42% S&P 500 \+ 36% MSCI Europe \+ 12% MSCI Japan \+ 10% CIMB Bank 1-Month Fixed Deposit Rate",
        r"MSCI ACWI Information Technology Index \+ MSCI ACWI Communication Services Index",
    ]:
        match = re.search(pattern, compact, re.I)
        if match:
            benchmark = match.group(0)
            break
    return volatility, fee, date, benchmark



def _ta_numeric_values(line):
    raw = line.strip().replace("—", "-").replace("–", "-")
    tokens = raw.replace("%", "").split()
    merged = []
    i = 0
    while i < len(tokens):
        token = tokens[i].rstrip(".;,:")
        if i + 1 < len(tokens) and re.fullmatch(r"[-+]?\d+\.\d", token) and re.fullmatch(r"\d", tokens[i + 1].rstrip(".;,:")):
            token = token + tokens[i + 1].rstrip(".;,:")
            i += 1
        elif i + 1 < len(tokens) and re.fullmatch(r"\d+", token) and re.fullmatch(r"\.\d+", tokens[i + 1].rstrip(".;,:")):
            token = token + tokens[i + 1].rstrip(".;,:")
            i += 1
        if token in {"-", "--", "N/A", "NA"}:
            merged.append(np.nan)
        elif re.fullmatch(r"[-+]?(?:\d+(?:\.\d*)?|\.\d+)", token):
            merged.append(float(token))
        i += 1
    return merged


def _ta_table_rows(lines, start, end=None, expected_rows=8, expected_values=5):
    end = len(lines) if end is None else end
    rows = []
    for line in lines[start:end]:
        values = _ta_numeric_values(line)
        if len(values) >= expected_values:
            rows.append(values[:expected_values])
            if len(rows) >= expected_rows:
                break
    return rows


def _parse_ta(lines, record):
    cumulative_start = _marker_index(lines, r"Cumulative Fund Performance") or 0
    annual_start = _marker_index(lines, r"^Annual Fund Performance")
    cumulative_rows = _ta_table_rows(lines, cumulative_start, annual_start, expected_rows=8, expected_values=8)
    if len(cumulative_rows) >= 8:
        fund = cumulative_rows[3]       # MYR Hedged row
        benchmark = cumulative_rows[7]  # TAGTF Benchmark row
        record["1Y Return (%)"] = fund[2]
        record["1Y Benchmark (%)"] = benchmark[2]
        _set_period(record, "3Y", fund_cum=fund[3], bench_cum=benchmark[3], basis="FFS reports cumulative 3Y; annualised equivalent calculated")
        _set_period(record, "5Y", fund_cum=fund[4], bench_cum=benchmark[4], basis="FFS reports cumulative 5Y; annualised equivalent calculated")
        record["YTD"] = (fund[5], benchmark[5])
    else:
        _append_warning(record, "TA MYR Hedged cumulative-performance rows were not detected")
    if annual_start is not None:
        years = [int(x) for x in re.findall(r"20\d{2}", lines[annual_start + 1])] if annual_start + 1 < len(lines) else [2025, 2024, 2023, 2022, 2021]
        annual_rows = _ta_table_rows(lines, annual_start, None, expected_rows=8, expected_values=len(years))
        if len(annual_rows) >= 8:
            fund_annual = annual_rows[3]
            benchmark_annual = annual_rows[7]
            record["Calendar"] = {
                year: (fund_annual[idx], benchmark_annual[idx])
                for idx, year in enumerate(years)
            }
    _append_warning(record, "TA volatility badge is not exposed cleanly in the text layer; verify manually")
    return record


def _parse_maybank(lines, record):
    total_start = _marker_index(lines, r"Total Return") or 0
    annual_header = _marker_index(lines, r"^Annual Management Fee") or len(lines)
    fund_total = _values_after_label(lines, r"^MGSEIF \(MYRH\)", start=total_start, stop=annual_header, count=8)
    bench_total = _values_after_label(lines, r"Benchmark", start=total_start, stop=annual_header, occurrence=2, count=8)
    fund_annual = _values_after_label(lines, r"^MGSEIF \(MYRH\)", start=annual_header, count=8)
    bench_annual = _values_after_label(lines, r"^Benchmark", start=annual_header, count=8)
    if len(fund_total) >= 7:
        record["1Y Return (%)"] = fund_total[4]
        record["1Y Benchmark (%)"] = bench_total[4] if len(bench_total) >= 5 else np.nan
        _set_period(record, "3Y", fund_cum=fund_total[5], fund_ann=fund_annual[0] if fund_annual else np.nan, bench_cum=bench_total[5] if len(bench_total) >= 6 else np.nan, bench_ann=bench_annual[0] if bench_annual else np.nan, basis="FFS reports cumulative and annualised 3Y")
        _set_period(record, "5Y", fund_cum=fund_total[6], fund_ann=fund_annual[1] if len(fund_annual) >= 2 else np.nan, bench_cum=bench_total[6] if len(bench_total) >= 7 else np.nan, bench_ann=bench_annual[1] if len(bench_annual) >= 2 else np.nan, basis="FFS reports cumulative and annualised 5Y")
        record["YTD"] = (fund_total[0], bench_total[0] if bench_total else np.nan)
    else:
        _append_warning(record, "Maybank MYRH total-return row was not detected")
    years = [int(x) for x in re.findall(r"20\d{2}", lines[annual_header])] if annual_header < len(lines) else []
    if len(years) >= 3 and len(fund_annual) >= 3:
        record["Calendar"] = {
            year: (fund_annual[idx + 3] if idx + 3 < len(fund_annual) else np.nan, bench_annual[idx + 3] if idx + 3 < len(bench_annual) else np.nan)
            for idx, year in enumerate(years)
        }
    return record


def _metadata(text, provider):
    compact = re.sub(r"\s+", " ", text)
    volatility = np.nan
    if provider != "TA":
        volatility = _first_float([
            r"Volatility Factor.*?for this Fund is\s+(\d+(?:\.\d+)?)",
            r"Volatility Factor.*?is\s+(\d+(?:\.\d+)?)",
        ], compact)
    fee = _first_float([
        r"Annual Management Fee.{0,100}?(\d+(?:\.\d+)?)\s*%",
        r"Management Fee.{0,100}?(\d+(?:\.\d+)?)\s*%",
    ], compact)
    if provider == "AHAM" and not np.isfinite(fee):
        fee = _first_float([r"Max\s+(\d+(?:\.\d+)?)\s*%\s*per annum"], compact)
    date = ""
    for pattern in [
        r"as (?:at|of)\s+(\d{1,2}[ /-][A-Za-z0-9]+[ /-]\d{2,4})",
        r"As of Date:\s*(\d{1,2}/\d{1,2}/\d{4})",
    ]:
        match = re.search(pattern, compact, re.I)
        if match:
            date = match.group(1)
            break
    benchmark = ""
    for pattern in [
        r"MSCI AC Asia Pacific ex Japan High Dividend Yield Index",
        r"BPAM Corporates All Bond Index",
        r"MSCI Emerging Markets Net Total Return Index",
        r"Dow Jones Islamic Market World Index",
        r"MSCI World / Health Care NR USD Index",
        r"42% S&P 500 \+ 36% MSCI Europe \+ 12% MSCI Japan \+ 10% CIMB Bank 1-Month Fixed Deposit Rate",
        r"MSCI ACWI Information Technology Index \+ MSCI ACWI Communication Services Index",
    ]:
        match = re.search(pattern, compact, re.I)
        if match:
            benchmark = match.group(0)
            break
    return volatility, fee, date, benchmark


# One final active override block. It is intentionally placed at the end of
# this module because earlier development iterations defined some helpers more
# than once while testing different PDF text modes.


def _parse_maybank(lines, record):
    total_start = _marker_index(lines, r"Total Return") or 0
    annual_header = _marker_index(lines, r"^Annual Management Fee") or len(lines)
    # The MYRH class-size line also contains the literal MYRH label. Exclude
    # that metadata line by requiring a numeric token immediately after the
    # class label.
    fund_total = _values_after_label(lines, r"^MGSEIF \(MYRH\)\s+[-+]?\d", start=total_start, stop=annual_header, count=8)
    bench_total = _values_after_label(lines, r"Benchmark\s+[-+]?\d", start=total_start, stop=annual_header, count=8)
    fund_annual = _values_after_label(lines, r"^MGSEIF \(MYRH\)\s+[-+]?\d", start=annual_header, count=8)
    bench_annual = _values_after_label(lines, r"Benchmark\s+[-+]?\d", start=annual_header, count=8)
    if len(fund_total) >= 7:
        record["1Y Return (%)"] = fund_total[4]
        record["1Y Benchmark (%)"] = bench_total[4] if len(bench_total) >= 5 else np.nan
        _set_period(record, "3Y", fund_cum=fund_total[5], fund_ann=fund_annual[0] if fund_annual else np.nan, bench_cum=bench_total[5] if len(bench_total) >= 6 else np.nan, bench_ann=bench_annual[0] if bench_annual else np.nan, basis="FFS reports cumulative and annualised 3Y")
        _set_period(record, "5Y", fund_cum=fund_total[6], fund_ann=fund_annual[1] if len(fund_annual) >= 2 else np.nan, bench_cum=bench_total[6] if len(bench_total) >= 7 else np.nan, bench_ann=bench_annual[1] if len(bench_annual) >= 2 else np.nan, basis="FFS reports cumulative and annualised 5Y")
        record["YTD"] = (fund_total[0], bench_total[0] if bench_total else np.nan)
    else:
        _append_warning(record, "Maybank MYRH total-return row was not detected")
    years = [int(x) for x in re.findall(r"20\d{2}", lines[annual_header])] if annual_header < len(lines) else []
    if len(years) >= 3 and len(fund_annual) >= 3:
        record["Calendar"] = {
            year: (fund_annual[idx + 3] if idx + 3 < len(fund_annual) else np.nan, bench_annual[idx + 3] if idx + 3 < len(bench_annual) else np.nan)
            for idx, year in enumerate(years)
        }
    return record


def _metadata(text, provider):
    compact = re.sub(r"\s+", " ", text)
    volatility = np.nan
    if provider != "TA":
        volatility = _first_float([
            r"Volatility Factor.*?for this Fund is\s+(\d+(?:\.\d+)?)",
            r"Volatility Factor.*?is\s+(\d+(?:\.\d+)?)",
            r"Fund Volatility\s+(\d+(?:\.\d+)?)",
        ], compact)
    fee = _first_float([
        r"Annual Management Fee.{0,100}?(\d+(?:\.\d+)?)\s*%",
        r"Management Fee.{0,100}?(\d+(?:\.\d+)?)\s*%",
    ], compact)
    if provider == "AHAM" and not np.isfinite(fee):
        fee = _first_float([r"Max\s+(\d+(?:\.\d+)?)\s*%\s*per annum"], compact)
    date = ""
    for pattern in [
        r"as (?:at|of)\s+(\d{1,2}[ /-][A-Za-z0-9]+[ /-]\d{2,4})",
        r"As of Date:\s*(\d{1,2}/\d{1,2}/\d{4})",
    ]:
        match = re.search(pattern, compact, re.I)
        if match:
            date = match.group(1)
            break
    benchmark = ""
    for pattern in [
        r"MSCI AC Asia Pacific ex Japan High Dividend Yield Index",
        r"BPAM Corporates All Bond Index",
        r"MSCI Emerging Markets Net Total Return Index",
        r"Dow Jones Islamic Market World Index",
        r"MSCI World / Health Care NR USD Index",
        r"42% S&P 500 \+ 36% MSCI Europe \+ 12% MSCI Japan \+ 10% CIMB Bank 1-Month Fixed Deposit Rate",
        r"MSCI ACWI Information Technology Index \+ MSCI ACWI Communication Services Index",
    ]:
        match = re.search(pattern, compact, re.I)
        if match:
            benchmark = match.group(0)
            break
    return volatility, fee, date, benchmark


def records_to_dataframe(records):
    rows = []
    years = sorted({year for record in records for year in record.get("Calendar", {})})
    ytd_years = set()
    for record in records:
        ytd = record.get("YTD", (np.nan, np.nan))
        date_match = re.search(r"(20\d{2})", str(record.get("As Of Date", "")))
        if date_match and any(np.isfinite(x) for x in ytd):
            ytd_years.add(int(date_match.group(1)))
    ytd_years = sorted(ytd_years)
    for record in records:
        row = {key: value for key, value in record.items() if key not in {"Calendar", "YTD"}}
        for year in years:
            fund, benchmark = record.get("Calendar", {}).get(year, (np.nan, np.nan))
            row[f"{year} Return (%)"] = fund
            row[f"{year} Benchmark (%)"] = benchmark
        as_of_year_match = re.search(r"(20\d{2})", str(record.get("As Of Date", "")))
        as_of_year = int(as_of_year_match.group(1)) if as_of_year_match else None
        fund_ytd, benchmark_ytd = record.get("YTD", (np.nan, np.nan))
        for year in ytd_years:
            row[f"{year} YTD Return (%)"] = fund_ytd if as_of_year == year else np.nan
            row[f"{year} YTD Benchmark (%)"] = benchmark_ytd if as_of_year == year else np.nan
        rows.append(row)
    df = pd.DataFrame(rows)
    preferred = [
        "Fund Name", "Provider", "Share Class", "As Of Date", "1Y Return (%)",
        "3Y Return (%)", "3Y Cumulative (%)", "3Y Annualised (%)", "3Y Return Source",
        "5Y Return (%)", "5Y Cumulative (%)", "5Y Annualised (%)", "5Y Return Source",
        "Volatility (%)", "Mgmt Fee (%)", "1Y Benchmark (%)", "3Y Benchmark (%)",
        "3Y Benchmark Cumulative (%)", "3Y Benchmark Annualised (%)", "5Y Benchmark (%)",
        "5Y Benchmark Cumulative (%)", "5Y Benchmark Annualised (%)", "Benchmark Name",
        "Return Basis", "Source File", "Extraction Warnings",
    ]
    dynamic = []
    for year in years:
        dynamic.extend([f"{year} Return (%)", f"{year} Benchmark (%)"])
    for year in ytd_years:
        dynamic.extend([f"{year} YTD Return (%)", f"{year} YTD Benchmark (%)"])
    ordered = [col for col in preferred + dynamic if col in df.columns]
    ordered += [col for col in df.columns if col not in ordered]
    return df[ordered]


# Small final fixes for the actual Maybank and AmBond text streams.

def _parse_maybank(lines, record):
    total_start = _marker_index(lines, r"Total Return") or 0
    annual_header = _marker_index(lines, r"^Annual Management Fee") or len(lines)
    fund_total = _values_after_label(lines, r"MGSEIF \(MYRH\)\s+[-+]?\d", start=total_start, stop=annual_header, count=8)
    bench_total = _values_after_label(lines, r"Benchmark\s+[-+]?\d", start=total_start, stop=annual_header, count=8)
    fund_annual = _values_after_label(lines, r"MGSEIF \(MYRH\)\s+[-+]?\d", start=annual_header, count=8)
    bench_annual = _values_after_label(lines, r"Benchmark\s+[-+]?\d", start=annual_header, count=8)
    if len(fund_total) >= 7:
        record["1Y Return (%)"] = fund_total[4]
        record["1Y Benchmark (%)"] = bench_total[4] if len(bench_total) >= 5 else np.nan
        _set_period(record, "3Y", fund_cum=fund_total[5], fund_ann=fund_annual[0] if fund_annual else np.nan, bench_cum=bench_total[5] if len(bench_total) >= 6 else np.nan, bench_ann=bench_annual[0] if bench_annual else np.nan, basis="FFS reports cumulative and annualised 3Y")
        _set_period(record, "5Y", fund_cum=fund_total[6], fund_ann=fund_annual[1] if len(fund_annual) >= 2 else np.nan, bench_cum=bench_total[6] if len(bench_total) >= 7 else np.nan, bench_ann=bench_annual[1] if len(bench_annual) >= 2 else np.nan, basis="FFS reports cumulative and annualised 5Y")
        record["YTD"] = (fund_total[0], bench_total[0] if bench_total else np.nan)
    else:
        _append_warning(record, "Maybank MYRH total-return row was not detected")
    years = [int(x) for x in re.findall(r"20\d{2}", lines[annual_header])] if annual_header < len(lines) else []
    if len(years) >= 3 and len(fund_annual) >= 3:
        record["Calendar"] = {
            year: (fund_annual[idx + 3] if idx + 3 < len(fund_annual) else np.nan, bench_annual[idx + 3] if idx + 3 < len(bench_annual) else np.nan)
            for idx, year in enumerate(years)
        }
    return record


def _metadata(text, provider):
    compact = re.sub(r"\s+", " ", text)
    volatility = np.nan
    if provider != "TA":
        volatility = _first_float([
            r"Volatility Factor.*?for this Fund is\s+(\d+(?:\.\d+)?)",
            r"Volatility Factor.*?is\s+(\d+(?:\.\d+)?)",
            r"Fund Volatility.*?(\d+(?:\.\d+)?)\s+(?:Very Low|Low|Moderate|High|Very High)",
        ], compact)
    fee = _first_float([
        r"Annual Management Fee.{0,100}?(\d+(?:\.\d+)?)\s*%",
        r"Management Fee.{0,100}?(\d+(?:\.\d+)?)\s*%",
    ], compact)
    if provider == "AHAM" and not np.isfinite(fee):
        fee = _first_float([r"Max\s+(\d+(?:\.\d+)?)\s*%\s*per annum"], compact)
    date = ""
    for pattern in [
        r"as (?:at|of)\s+(\d{1,2}[ /-][A-Za-z0-9]+[ /-]\d{2,4})",
        r"As of Date:\s*(\d{1,2}/\d{1,2}/\d{4})",
    ]:
        match = re.search(pattern, compact, re.I)
        if match:
            date = match.group(1)
            break
    benchmark = ""
    for pattern in [
        r"MSCI AC Asia Pacific ex Japan High Dividend Yield Index",
        r"BPAM Corporates All Bond Index",
        r"MSCI Emerging Markets Net Total Return Index",
        r"Dow Jones Islamic Market World Index",
        r"MSCI World / Health Care NR USD Index",
        r"42% S&P 500 \+ 36% MSCI Europe \+ 12% MSCI Japan \+ 10% CIMB Bank 1-Month Fixed Deposit Rate",
        r"MSCI ACWI Information Technology Index \+ MSCI ACWI Communication Services Index",
    ]:
        match = re.search(pattern, compact, re.I)
        if match:
            benchmark = match.group(0)
            break
    return volatility, fee, date, benchmark



def _parse_maybank(lines, record):
    total_start = _marker_index(lines, r"Total Return") or 0
    annual_header = _marker_index(lines, r"^Annual Management Fee") or len(lines)
    fund_label = r"MGSEIF \(MYRH\)(?=\s+[-+]?\d)"
    benchmark_label = r"Benchmark(?=\s+[-+]?\d)"
    fund_total = _values_after_label(lines, fund_label, start=total_start, stop=annual_header, count=8)
    bench_total = _values_after_label(lines, benchmark_label, start=total_start, stop=annual_header, count=8)
    fund_annual = _values_after_label(lines, fund_label, start=annual_header, count=8)
    bench_annual = _values_after_label(lines, benchmark_label, start=annual_header, count=8)
    if len(fund_total) >= 7:
        record["1Y Return (%)"] = fund_total[4]
        record["1Y Benchmark (%)"] = bench_total[4] if len(bench_total) >= 5 else np.nan
        _set_period(record, "3Y", fund_cum=fund_total[5], fund_ann=fund_annual[0] if fund_annual else np.nan, bench_cum=bench_total[5] if len(bench_total) >= 6 else np.nan, bench_ann=bench_annual[0] if bench_annual else np.nan, basis="FFS reports cumulative and annualised 3Y")
        _set_period(record, "5Y", fund_cum=fund_total[6], fund_ann=fund_annual[1] if len(fund_annual) >= 2 else np.nan, bench_cum=bench_total[6] if len(bench_total) >= 7 else np.nan, bench_ann=bench_annual[1] if len(bench_annual) >= 2 else np.nan, basis="FFS reports cumulative and annualised 5Y")
        record["YTD"] = (fund_total[0], bench_total[0] if bench_total else np.nan)
    else:
        _append_warning(record, "Maybank MYRH total-return row was not detected")
    years = [int(x) for x in re.findall(r"20\d{2}", lines[annual_header])] if annual_header < len(lines) else []
    if len(years) >= 3 and len(fund_annual) >= 3:
        record["Calendar"] = {
            year: (fund_annual[idx + 3] if idx + 3 < len(fund_annual) else np.nan, bench_annual[idx + 3] if idx + 3 < len(bench_annual) else np.nan)
            for idx, year in enumerate(years)
        }
    return record


# ============================================================
# SECTION: EMAIL & FIREBASE FUNCTIONS
# ============================================================
def send_otp_email(recipient_email, otp_code):
    try:
        message = MIMEMultipart()
        message["From"] = GMAIL_ADDRESS
        message["To"] = recipient_email
        message["Subject"] = "Your Portfolio Analyzer OTP Code"
        body = f"Hello,\n\nYour OTP for the Portfolio Analyzer Tool is: {otp_code}\n\nBest regards,\nChew Advisory"
        message.attach(MIMEText(body, "plain"))
        server = smtplib.SMTP_SSL("smtp.gmail.com", 465)
        server.login(GMAIL_ADDRESS, GMAIL_APP_PASSWORD)
        server.sendmail(GMAIL_ADDRESS, recipient_email, message.as_string())
        server.quit()
        return True
    except Exception as e:
        st.error(f"Error sending OTP: {str(e)}")
        return False

def get_user_stats(email):
    if email == ADMIN_EMAIL: return "allowed", 0, 0, 999999
    if not db: return "allowed", 0, 0, 3 
    try:
        users_ref = db.collection('user_usage').where('email', '==', email).limit(1).get()
        docs = list(users_ref)
        if docs:
            data = docs[0].to_dict()
            if data.get('deleted_at') is not None: return "deleted", 0, 0, 0
            return "allowed", int(data.get('access_count', 0)), int(data.get('generation_count', 0)), int(data.get('max_limit', 3))
        else:
            db.collection('user_usage').add({'email': email, 'access_count': 0, 'generation_count': 0, 'max_limit': 3, 'created_at': firestore.SERVER_TIMESTAMP})
            return "allowed", 0, 0, 3
    except: return "allowed", 0, 0, 3

def check_access_allowed(email):
    status, acc, gen, lim = get_user_stats(email)
    if status != "allowed" or acc >= lim or gen >= lim: return False, lim, acc, gen
    return True, lim, acc, gen

def increment_access(email):
    if not db or email == ADMIN_EMAIL: return
    try:
        users_ref = db.collection('user_usage').where('email', '==', email).limit(1).get()
        docs = list(users_ref)
        if docs:
            doc_id = docs[0].id
            current = int(docs[0].to_dict().get('access_count', 0))
            db.collection('user_usage').document(doc_id).update({'access_count': current + 1, 'last_accessed_at': firestore.SERVER_TIMESTAMP})
    except: pass

def increment_generation(email):
    if not db or email == ADMIN_EMAIL: return
    try:
        users_ref = db.collection('user_usage').where('email', '==', email).limit(1).get()
        docs = list(users_ref)
        if docs:
            doc_id = docs[0].id
            current = int(docs[0].to_dict().get('generation_count', 0))
            db.collection('user_usage').document(doc_id).update({'generation_count': current + 1})
    except: pass

# ============================================================
# SECTION: PORTFOLIO CALCULATIONS & OPTIMIZATION
# ============================================================
def calculate_required_cagr(target_sum, initial_investment, monthly_contribution, years):
    if years <= 0 or target_sum <= 0: return 0.0
    r = 0.05 
    for _ in range(50): 
        fv_guess = initial_investment * ((1 + r) ** years) + (monthly_contribution * 12 * (((1 + r) ** years - 1) / r) if r > 0 else monthly_contribution * 12 * years)
        derivative = initial_investment * years * ((1 + r) ** (years - 1))
        if r > 0: derivative += monthly_contribution * 12 * (years * (1 + r)**(years - 1) * r - ((1 + r)**years - 1)) / (r**2)
        if abs(derivative) < 1e-8: break
        r = r - (fv_guess - target_sum) / derivative
        if r < -0.5: r = -0.5 
    return max(0.0, r * 100)

def calculate_future_value(initial_investment, monthly_contribution, years, annual_return_pct):
    r = annual_return_pct / 100
    if r == 0:
        fv = initial_investment + (monthly_contribution * 12 * years)
    else:
        fv = initial_investment * ((1 + r) ** years) + (monthly_contribution * 12 * (((1 + r) ** years - 1) / r))
    return fv

def optimize_portfolio_max_return(df, risk_profile):
    """
    Optimizes portfolio to MAXIMIZE return while strictly staying within the 
    max allowable volatility for the risk profile. Includes ALL funds (min 5%).
    """
    n = len(df)
    if n == 0: return np.array([]), 0, 0
    
    risk_thresholds = {"Conservative": 10.0, "Moderate": 15.0, "Growth": 20.0}
    max_volatility = risk_thresholds.get(risk_profile, 15.0)
    
    returns_col = _analysis_return_series(df)
    returns = returns_col.fillna(0).values
    volatilities = pd.to_numeric(df['Volatility (%)'], errors='coerce').values
    
    # Start with minimum 5% for all funds to ensure all are included
    weights = np.full(n, 0.05)
    remaining_weight = 1.0 - np.sum(weights)
    
    # Sort indices by return (highest first)
    sorted_indices = np.argsort(returns)[::-1]
    
    # Greedy allocation: Add weight to highest returning funds first
    for idx in sorted_indices:
        if remaining_weight <= 0.001:
            break
            
        # Max weight for a single fund is 40%
        max_add = 0.40 - weights[idx]
        add_weight = min(max_add, remaining_weight)
        
        # Step-down search to find max weight that fits volatility limit
        step = add_weight
        while step > 0.001:
            test_weights = weights.copy()
            test_weights[idx] += step
            test_vol = np.sum(volatilities * test_weights)
            
            if test_vol <= max_volatility:
                weights[idx] += step
                remaining_weight -= step
                break
            else:
                step /= 2.0
                
    # Normalize to ensure exactly 100%
    weights = weights / np.sum(weights)
    
    opt_return = np.sum(returns * weights)
    opt_vol = np.sum(volatilities * weights)
    
    return weights, opt_return, opt_vol

def calculate_portfolio_metrics(df, weights, year_returns_cols):
    returns_1y = _analysis_return_series(df)
    volatilities = df['Volatility (%)']
    fees = df['Mgmt Fee (%)']
    
    valid_return_mask = returns_1y.notna()
    valid_vol_mask = volatilities.notna()
    valid_fee_mask = fees.notna()
    
    portfolio_return = np.sum(returns_1y[valid_return_mask] * weights[valid_return_mask]) if valid_return_mask.any() else 0
    portfolio_volatility = np.sum(volatilities[valid_vol_mask] * weights[valid_vol_mask]) if valid_vol_mask.any() else 0
    portfolio_fee = np.sum(fees[valid_fee_mask] * weights[valid_fee_mask]) if valid_fee_mask.any() else 0
    
    risk_adjusted = portfolio_return / portfolio_volatility if portfolio_volatility > 0 else 0
    
    yearly_returns = []
    for year, col_name in year_returns_cols:
        if col_name in df.columns:
            year_data = df[col_name]
            valid_mask = year_data.notna()
            if valid_mask.any():
                year_return = np.sum(year_data[valid_mask] * weights[valid_mask])
                yearly_returns.append(year_return)
    
    best_year = max(yearly_returns) if yearly_returns else None
    worst_year = min(yearly_returns) if yearly_returns else None
    avg_yearly = np.mean(yearly_returns) if yearly_returns else None
    positive_years = sum(1 for r in yearly_returns if r > 0)
    consistency = (positive_years / len(yearly_returns) * 100) if yearly_returns else 0
    
    return {
        'return': portfolio_return,
        'volatility': portfolio_volatility,
        'fee': portfolio_fee,
        'risk_adjusted': risk_adjusted,
        'best_year': best_year,
        'worst_year': worst_year,
        'avg_yearly': avg_yearly,
        'consistency': consistency,
        'yearly_returns': yearly_returns
    }

# ============================================================
# SECTION: LOGIN PAGE
# ============================================================
def show_login_page():
    st.markdown("<h2 style='text-align: center;'>🔐 Login to Portfolio Analyzer</h2>", unsafe_allow_html=True)
    if firebase_init_error:
        st.warning("Firebase usage tracking is unavailable in this deployment. The login page remains available, but Firebase-based usage limits will not be applied until FIREBASE_SERVICE_ACCOUNT is configured in Streamlit Secrets.")
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        email = st.text_input("Enter your email address:", placeholder="your.email@example.com", key="login_email")
        if st.button("Send OTP Code", use_container_width=True, key="send_otp_btn"):
            if "@" in email and "." in email.split("@")[1]:
                otp = generate_otp()
                if send_otp_email(email, otp):
                    st.session_state.otp_email = email
                    st.session_state.otp_code = otp
                    st.session_state.show_otp_input = True
                    st.success(f"✅ OTP sent to {email}. Check your email!")
                else: st.error("Failed to send OTP.")
            else: st.error("Please enter a valid email address.")
        
        if st.session_state.get('show_otp_input', False):
            st.info("An OTP code has been sent to your email.")
            otp_input = st.text_input("Enter 6-digit OTP:", placeholder="000000", key="otp_input", type="password")
            if st.button("Verify OTP", use_container_width=True, key="verify_otp_btn"):
                if otp_input == st.session_state.otp_code:
                    allowed, lim, acc, gen = check_access_allowed(email)
                    if allowed:
                        st.session_state.authenticated = True
                        st.session_state.user_email = email
                        increment_access(email)
                        st.success("✅ Login successful!")
                        st.rerun()
                    else: st.error(f"❌ Limit reached ({lim}). Contact cktchew@gmail.com.")
                else: st.error("❌ Incorrect OTP.")

# ============================================================
# SECTION: MAIN APP LOGIC
# ============================================================
if not st.session_state.authenticated:
    show_login_page()
else:
    st.markdown("""
    <style>
    .main-header { text-align: center; color: #1f77b4; font-size: 2.5em; font-weight: bold; margin-bottom: 10px; }
    .sub-header { text-align: center; color: #666; font-size: 1.1em; margin-bottom: 20px; }
    </style>
    <div class="main-header">CHEW ADVISORY</div>
    <div class="sub-header">Portfolio Analysis Tool</div>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns([0.9, 0.1])
    with col2:
        if st.button("Logout", key="btn_logout"):
            for key in ['authenticated', 'user_email', 'user_name', 'page', 'portfolio_data', 'show_otp_input', 'otp_code', 'otp_email', 'funds_df']:
                st.session_state[key] = False if key == 'authenticated' else (None if key in ['user_email', 'user_name', 'otp_code', 'otp_email'] else 'home' if key == 'page' else None if key == 'funds_df' else {})
            st.rerun()

    if st.session_state.page == 'home':
        st.markdown("---")
        st.markdown("### Welcome to the Portfolio Analyzer")
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("📝 Start New Analysis", use_container_width=True, type="primary"):
                st.session_state.page = 'input'
                st.session_state.portfolio_data = {}
                st.session_state.funds_df = None
                st.session_state.user_name = None
                st.rerun()
        with col2:
            if st.session_state.user_email == ADMIN_EMAIL: st.info(" Admin: Unlimited access")

    elif st.session_state.page == 'input':
        st.header("Step 1: Define Goal & Upload Fund Data")

        if st.session_state.portfolio_data:
            st.info("You have existing analysis results from your previous upload.")
            if st.button("Return to Previous Analysis", use_container_width=True, type="primary", key="btn_return_analysis"):
                st.session_state.page = 'analysis'
                st.rerun()
            st.markdown("---")
            if st.button("Clear Data & Start New Analysis", use_container_width=True, key="btn_clear_data"):
                st.session_state.portfolio_data = {}
                st.session_state.funds_df = None
                st.session_state.pop("ffs_parsed_df", None)
                st.session_state.pop("ffs_upload_signature", None)
                st.rerun()
            st.markdown("---")

        st.subheader("Client Information")
        client_name = st.text_input("Client Name (Optional):", placeholder="Enter your name", value=st.session_state.get('user_name', '') or '')
        st.session_state.user_name = client_name if client_name else None

        st.subheader("1. Investment Goal & Risk Profile")
        col1, col2, col3 = st.columns(3)
        with col1:
            goal_type = st.radio("Goal Type:", ["Reach a Target Sum ($)", "Achieve Target Annual Growth (%)"])
        with col2:
            years = st.number_input("Time Horizon (Years)", min_value=1, max_value=50, value=10)
        with col3:
            risk_profile = st.selectbox("Risk Profile:", ["Conservative", "Moderate", "Growth"])

        if goal_type == "Reach a Target Sum ($)":
            target_value = st.number_input("Target Final Sum ($)", min_value=1000, value=100000, step=1000)
            target_growth = None
        else:
            target_growth = st.number_input("Target Annual Growth (%)", min_value=1.0, max_value=20.0, value=8.0, step=0.5)
            target_value = None

        st.subheader("2. Capital & Contributions")
        col1, col2 = st.columns(2)
        with col1:
            initial_investment = st.number_input("Initial Lump Sum Investment ($)", min_value=0, value=10000, step=1000)
        with col2:
            monthly_contribution = st.number_input("Monthly Contribution ($)", min_value=0, value=500, step=100)

        st.subheader("3. Upload Fund Data")
        st.info(
            "Upload selectable-text Fund Fact Sheets. The app extracts the data locally, "
            "shows the extracted table for correction, and does not analyse anything until "
            "you click the analysis button. Scanned image-only PDFs are flagged because OCR "
            "is intentionally not used in this version."
        )

        input_mode = st.radio(
            "Input source",
            ["Fund Fact Sheets (PDF)", "Legacy Excel file"],
            horizontal=True,
            key="input_mode",
        )

        template_columns = [
            "Fund Name", "1Y Return (%)", "3Y Cumulative (%)", "3Y Annualised (%)", "3Y Return (%)",
            "5Y Cumulative (%)", "5Y Annualised (%)", "5Y Return (%)", "Volatility (%)", "Mgmt Fee (%)",
            "1Y Benchmark (%)", "3Y Benchmark Cumulative (%)", "3Y Benchmark Annualised (%)",
            "5Y Benchmark Cumulative (%)", "5Y Benchmark Annualised (%)", "Benchmark Name",
            "2021 Return (%)", "2022 Return (%)", "2023 Return (%)", "2024 Return (%)", "2025 Return (%)",
            "2021 Benchmark (%)", "2022 Benchmark (%)", "2023 Benchmark (%)", "2024 Benchmark (%)", "2025 Benchmark (%)",
        ]

        if input_mode == "Fund Fact Sheets (PDF)":
            uploaded_files = st.file_uploader(
                "Upload one or more Fund Fact Sheets (PDF)",
                type=["pdf"],
                accept_multiple_files=True,
                key="ffs_pdf_upload",
            )
            if uploaded_files:
                signature_text = "|".join(f"{file.name}:{file.size}" for file in uploaded_files)
                upload_signature = hashlib.sha256(signature_text.encode("utf-8")).hexdigest()
                if st.session_state.get("ffs_upload_signature") != upload_signature:
                    records = []
                    parse_errors = []
                    progress = st.progress(0, text="Extracting Fund Fact Sheets...")
                    for index, uploaded in enumerate(uploaded_files, start=1):
                        try:
                            records.append(parse_ffs_bytes(uploaded.getvalue(), uploaded.name))
                        except Exception as exc:
                            parse_errors.append({"Source File": uploaded.name, "Error": str(exc)})
                        progress.progress(index / len(uploaded_files), text=f"Processed {index} of {len(uploaded_files)} PDF(s)")
                    progress.empty()
                    st.session_state.ffs_upload_signature = upload_signature
                    st.session_state.ffs_parse_errors = parse_errors
                    st.session_state.ffs_parsed_df = records_to_dataframe(records)

                if st.session_state.get("ffs_parse_errors"):
                    st.warning("Some PDFs could not be parsed. They are listed below and can be re-uploaded as selectable-text PDFs.")
                    st.dataframe(pd.DataFrame(st.session_state.ffs_parse_errors), use_container_width=True, hide_index=True)

                extracted_df = st.session_state.get("ffs_parsed_df", pd.DataFrame())
                if extracted_df.empty:
                    st.error("No usable FFS records were extracted.")
                else:
                    st.markdown("### Review and correct extracted data")
                    st.caption(
                        "Edit any cell that needs correction. The 3Y and 5Y cumulative and annualised columns are shown separately. "
                        "The legacy 3Y Return (%) and 5Y Return (%) columns are annualised aliases used by the analysis."
                    )
                    editor_key = f"ffs_editor_{st.session_state.get('ffs_upload_signature', 'empty')[:12]}"
                    edited_df = st.data_editor(
                        extracted_df,
                        key=editor_key,
                        num_rows="dynamic",
                        use_container_width=True,
                        hide_index=True,
                    )
                    st.session_state.ffs_current_edit_df = edited_df.copy()
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
                        edited_df.to_excel(writer, sheet_name="Fund Data", index=False)
                    st.download_button(
                        "Download Extracted Data (Excel)",
                        data=excel_buffer.getvalue(),
                        file_name="fund_data_extracted_from_ffs.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="download_ffs_data",
                    )

                    st.markdown("---")
                    if st.button("Analyse and Generate Report/Charts", use_container_width=True, type="primary", key="btn_calc_ffs_analysis"):
                        df = _normalize_analysis_columns(edited_df)
                        required_cols = ["Fund Name", "1Y Return (%)", "Volatility (%)", "Mgmt Fee (%)"]
                        missing_cols = [col for col in required_cols if col not in df.columns]
                        missing_values = [col for col in required_cols if col in df.columns and df[col].isna().any()]
                        if missing_cols:
                            st.error(f"Missing required columns: {', '.join(missing_cols)}")
                        elif missing_values:
                            st.error(
                                "Please correct the highlighted missing required values before analysis: "
                                + ", ".join(missing_values)
                                + ". TA Global Technology volatility, for example, may require manual entry because its badge is not exposed in the PDF text layer."
                            )
                        elif len(df) == 0:
                            st.error("At least one fund is required.")
                        else:
                            year_returns, year_benchmarks = detect_year_columns(df)
                            st.session_state.funds_df = df
                            st.session_state.portfolio_data = {
                                'client_name': st.session_state.user_name,
                                'goal_type': goal_type, 'target_value': target_value, 'target_growth': target_growth,
                                'years': years, 'initial_investment': initial_investment, 'monthly_contribution': monthly_contribution,
                                'risk_profile': risk_profile, 'funds_df': df,
                                'year_returns': year_returns, 'year_benchmarks': year_benchmarks,
                                'input_source': 'Fund Fact Sheets (PDF)',
                            }
                            increment_generation(st.session_state.user_email)
                            st.session_state.page = 'analysis'
                            st.rerun()
            else:
                st.warning("Please upload at least one Fund Fact Sheet PDF to proceed.")

        else:
            st.info("The legacy Excel option is retained for backward compatibility. The FFS workflow above is the recommended input method.")
            if st.button("Download Excel Template", key="btn_dl_template"):
                template_df = pd.DataFrame(columns=template_columns)
                buffer = io.BytesIO()
                with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                    template_df.to_excel(writer, sheet_name='Fund Data', index=False)
                st.download_button(
                    label="Download Template Excel",
                    data=buffer.getvalue(),
                    file_name="fund_data_template.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_excel_template",
                )

            uploaded_excel = st.file_uploader("Upload Excel File", type=['xlsx', 'xls'], key="legacy_excel_upload")
            if uploaded_excel:
                try:
                    df = pd.read_excel(uploaded_excel)
                    df.columns = df.columns.astype(str).str.strip()
                    df = _normalize_analysis_columns(df)
                    required_cols = ['Fund Name', '1Y Return (%)', 'Volatility (%)', 'Mgmt Fee (%)']
                    missing_cols = [col for col in required_cols if col not in df.columns]
                    if missing_cols:
                        st.error(f"Missing required columns: {', '.join(missing_cols)}")
                    else:
                        st.markdown("### Review and correct fund data")
                        edited_df = st.data_editor(df, key="legacy_excel_editor", num_rows="dynamic", use_container_width=True, hide_index=True)
                        st.session_state.funds_df = edited_df.copy()
                        if st.button("Analyse and Generate Report/Charts", use_container_width=True, type="primary", key="btn_calc_excel_analysis"):
                            final_df = _normalize_analysis_columns(edited_df)
                            missing_values = [col for col in required_cols if final_df[col].isna().any()]
                            if missing_values:
                                st.error("Please correct missing required values: " + ", ".join(missing_values))
                            else:
                                year_returns, year_benchmarks = detect_year_columns(final_df)
                                st.session_state.portfolio_data = {
                                    'client_name': st.session_state.user_name,
                                    'goal_type': goal_type, 'target_value': target_value, 'target_growth': target_growth,
                                    'years': years, 'initial_investment': initial_investment, 'monthly_contribution': monthly_contribution,
                                    'risk_profile': risk_profile, 'funds_df': final_df,
                                    'year_returns': year_returns, 'year_benchmarks': year_benchmarks,
                                    'input_source': 'Legacy Excel file',
                                }
                                increment_generation(st.session_state.user_email)
                                st.session_state.page = 'analysis'
                                st.rerun()
                except Exception as exc:
                    st.error(f"Error reading Excel file: {exc}")
            else:
                st.warning("Please upload an Excel file to proceed.")

        if st.button("Back to Home", key="btn_back_home_input"):
            st.session_state.page = 'home'
            st.rerun()

    elif st.session_state.page == 'analysis':
        st.header("📊 Step 2: Portfolio Analysis")
        data = st.session_state.portfolio_data
        df = data['funds_df'].copy()
        year_returns_cols = data['year_returns']
        year_benchmarks_cols = data['year_benchmarks']
        
        if data['goal_type'] == "Reach a Target Sum ($)":
            target_return = calculate_required_cagr(data['target_value'], data['initial_investment'], data['monthly_contribution'], data['years'])
            target_amount = data['target_value']
        else:
            target_return = data['target_growth']
            target_amount = calculate_future_value(data['initial_investment'], data['monthly_contribution'], data['years'], target_return)

        n = len(df)
        equal_weights = np.ones(n) / n
        eq_metrics = calculate_portfolio_metrics(df, equal_weights, year_returns_cols)
        eq_amount = calculate_future_value(data['initial_investment'], data['monthly_contribution'], data['years'], eq_metrics['return'])
        
        # Optimize for MAX RETURN within risk profile
        opt_weights, opt_return, opt_vol = optimize_portfolio_max_return(df, data['risk_profile'])
        opt_metrics = calculate_portfolio_metrics(df, opt_weights, year_returns_cols)
        opt_amount = calculate_future_value(data['initial_investment'], data['monthly_contribution'], data['years'], opt_metrics['return'])

        risk_thresholds = {"Conservative": 10.0, "Moderate": 15.0, "Growth": 20.0}
        max_vol_threshold = risk_thresholds[data['risk_profile']]

        st.subheader("Portfolio Performance Analysis")
        col1, col2, col3 = st.columns(3)
        with col1: st.metric("Equal-Weighted Return", f"{eq_metrics['return']:.1f}% p.a.")
        with col2: st.metric("Optimized Return (Max)", f"{opt_metrics['return']:.1f}% p.a.")
        with col3: st.metric("Target Return", f"{target_return:.1f}% p.a.")

        st.subheader("Goal Feasibility & Risk Assessment")
        st.info(f"**Max Allowable Volatility for {data['risk_profile']} Profile:** {max_vol_threshold}%")
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Equal-Weighted Portfolio", "✅ Achievable" if eq_amount >= target_amount else "️ Shortfall", delta=f"${eq_amount - target_amount:,.0f}")
            st.caption(f"Volatility: {eq_metrics['volatility']:.1f}% ({'✅ Matches' if eq_metrics['volatility'] <= max_vol_threshold else '⚠️ Exceeds'} {data['risk_profile']} profile)")
        with col2:
            st.metric("Optimized Portfolio", "✅ Achievable" if opt_amount >= target_amount else "⚠️ Shortfall", delta=f"${opt_amount - target_amount:,.0f}")
            st.caption(f"Volatility: {opt_metrics['volatility']:.1f}% ({'✅ Matches' if opt_metrics['volatility'] <= max_vol_threshold else '⚠️ Exceeds'} {data['risk_profile']} profile)")

        # Helper functions for recommendations
        def find_monthly_for_target(target, initial, years, return_rate):
            r = return_rate / 100
            if r == 0: return max(0, (target - initial) / (years * 12))
            fv_initial = initial * ((1 + r) ** years)
            remaining = target - fv_initial
            if remaining <= 0: return 0
            return max(0, remaining * r / (12 * (((1 + r) ** years) - 1)))
        
        def find_initial_for_target(target, monthly, years, return_rate):
            r = return_rate / 100
            if r == 0: return max(0, target - (monthly * 12 * years))
            fv_monthly = monthly * 12 * (((1 + r) ** years) - 1) / r
            remaining = target - fv_monthly
            if remaining <= 0: return 0
            return max(0, remaining / ((1 + r) ** years))
        
        def find_years_for_target(target, initial, monthly, return_rate):
            for test_years in range(1, 100):
                if calculate_future_value(initial, monthly, test_years, return_rate) >= target:
                    return test_years
            return 100

        # Calculate recommendations based on OPTIMIZED portfolio max return
        req_monthly = find_monthly_for_target(target_amount, data['initial_investment'], data['years'], opt_metrics['return'])
        req_initial_zero_monthly = find_initial_for_target(target_amount, 0, data['years'], opt_metrics['return'])
        req_years = find_years_for_target(target_amount, data['initial_investment'], data['monthly_contribution'], opt_metrics['return'])

        st.subheader("💡 Recommendations (Based on Optimized Portfolio)")
        
        # Option 1: Monthly Contribution
        diff_monthly = req_monthly - data['monthly_contribution']
        if diff_monthly < -1:
            msg1 = f"You can **reduce** your monthly contribution by **${abs(diff_monthly):,.0f}** (New total: ${req_monthly:,.0f}/month) and still meet your target."
        elif diff_monthly > 1:
            msg1 = f"You need to **increase** your monthly contribution by **${diff_monthly:,.0f}** (New total: ${req_monthly:,.0f}/month) to meet your target."
        else:
            msg1 = "Your current monthly contribution is exactly on track to meet your target."

        # Option 2: Initial Capital (Assuming $0 monthly)
        diff_initial = req_initial_zero_monthly - data['initial_investment']
        if diff_initial < -1:
            msg2 = f"If you stop monthly contributions ($0/month), you can **reduce** your initial investment by **${abs(diff_initial):,.0f}** (New total: ${req_initial_zero_monthly:,.0f}) and still meet your target."
        elif diff_initial > 1:
            msg2 = f"If you stop monthly contributions ($0/month), you need to **increase** your initial investment by **${diff_initial:,.0f}** (New total: ${req_initial_zero_monthly:,.0f}) to meet your target."
        else:
            msg2 = "If you stop monthly contributions ($0/month), your current initial investment is exactly on track."

        # Option 3: Duration
        diff_years = req_years - data['years']
        if diff_years < 0:
            msg3 = f"You can **reduce** your time horizon by **{abs(diff_years)} years** (New total: {req_years} years) and still meet your target."
        elif diff_years > 0:
            msg3 = f"You need to **increase** your time horizon by **{diff_years} years** (New total: {req_years} years) to meet your target."
        else:
            msg3 = "Your current time horizon is exactly on track to meet your target."

        st.info(f"**Option 1 - Adjust Monthly Contribution:**\n{msg1}")
        st.info(f"**Option 2 - Adjust Initial Capital (with $0 monthly):**\n{msg2}")
        st.info(f"**Option 3 - Adjust Time Horizon:**\n{msg3}")

        # Charts Section
        st.subheader("📊 Portfolio Allocation")
        col1, col2 = st.columns(2)
        with col1:
            fig1, ax1 = plt.subplots(figsize=(8, 8))
            fig1.suptitle("Portfolio Allocation", fontsize=16, fontweight='bold', y=0.98)
            wedges1, texts1, autotexts1 = ax1.pie(equal_weights, labels=df['Fund Name'], autopct='%1.1f%%', startangle=90, colors=plt.cm.Paired.colors, pctdistance=0.85, labeldistance=1.1)
            for autotext in autotexts1: autotext.set_color('white'); autotext.set_fontweight('bold'); autotext.set_fontsize(9)
            for text in texts1: text.set_fontsize(8)
            ax1.set_title("Equal-Weighted Allocation", fontsize=12, fontweight='bold', pad=20)
            fig1.tight_layout(rect=[0, 0, 1, 0.93])
            st.pyplot(fig1)
        
        with col2:
            fig2, ax2 = plt.subplots(figsize=(8, 8))
            wedges2, texts2, autotexts2 = ax2.pie(opt_weights, labels=df['Fund Name'], autopct='%1.1f%%', startangle=90, colors=plt.cm.Paired.colors, pctdistance=0.85, labeldistance=1.1)
            for autotext in autotexts2: autotext.set_color('white'); autotext.set_fontweight('bold'); autotext.set_fontsize(9)
            for text in texts2: text.set_fontsize(8)
            ax2.set_title("Optimized Allocation (Max Return)", fontsize=12, fontweight='bold', pad=20)
            st.pyplot(fig2)

        st.subheader("📈 Portfolio Performance vs Benchmark")
        years_list = sorted(list(set([year for year, col in year_returns_cols])))
        fig3, ax3 = plt.subplots(figsize=(12, 7))
        has_data = False; has_benchmark = False
        eq_port_returns = []; opt_port_returns = []; bench_returns = []
        benchmark_names = df['Benchmark Name'].dropna().astype(str).str.strip().unique() if 'Benchmark Name' in df.columns else []
        benchmark_label = str(benchmark_names[0]) if len(benchmark_names) == 1 else "Composite Benchmark"
        
        for year in years_list:
            year_return_col = f'{year} Return (%)'
            year_bench_col = f'{year} Benchmark (%)'
            if year_return_col in df.columns:
                valid_mask = df[year_return_col].notna()
                if valid_mask.any():
                    eq_port_returns.append(np.sum(df[year_return_col][valid_mask] * equal_weights[valid_mask]))
                    opt_port_returns.append(np.sum(df[year_return_col][valid_mask] * opt_weights[valid_mask]))
                    has_data = True
                else: eq_port_returns.append(np.nan); opt_port_returns.append(np.nan)
            else: eq_port_returns.append(np.nan); opt_port_returns.append(np.nan)
            
            if year_bench_col in df.columns:
                valid_mask = df[year_bench_col].notna()
                if valid_mask.any():
                    bench_returns.append(np.sum(df[year_bench_col][valid_mask] * equal_weights[valid_mask]))
                    has_benchmark = True
                else: bench_returns.append(np.nan)
            else: bench_returns.append(np.nan)
        
        if has_data:
            ax3.plot(years_list, eq_port_returns, marker='o', linewidth=3, label='Equal-Weighted Portfolio', color='#3498db', markersize=10)
            ax3.plot(years_list, opt_port_returns, marker='s', linewidth=3, label='Optimized Portfolio', color='#2ecc71', markersize=10)
        if has_benchmark:
            ax3.plot(years_list, bench_returns, marker='^', linewidth=3, label=benchmark_label, color='#e74c3c', markersize=10)
        
        if has_data or has_benchmark:
            ax3.set_xlabel("Year", fontsize=12, fontweight='bold'); ax3.set_ylabel("Calendar-Year Return (%)", fontsize=12, fontweight='bold')
            ax3.set_title(f"Portfolio Calendar-Year Performance vs {benchmark_label}", fontsize=14, fontweight='bold')
            ax3.legend(loc='upper left', fontsize=11); ax3.grid(True, alpha=0.3); ax3.axhline(y=0, color='black', linestyle='-', linewidth=0.5)
            st.pyplot(fig3)
        else: st.warning("No calendar year return data found.")

        st.subheader("📉 Individual Fund Performance")
        fig4, ax4 = plt.subplots(figsize=(14, 8))
        fund_has_data = False; colors = plt.cm.tab10.colors; full_years = [y for y, c in year_returns_cols]
        for i, (_, row) in enumerate(df.iterrows()):
            fund_returns = []
            for year, col_name in year_returns_cols:
                if col_name in df.columns and pd.notna(row.get(col_name)): fund_returns.append(row[col_name])
                else: fund_returns.append(np.nan)
            if any(pd.notna(r) for r in fund_returns):
                ax4.plot(full_years, fund_returns, marker='o', linewidth=2, label=row['Fund Name'][:30], color=colors[i % len(colors)], markersize=8, alpha=0.8)
                fund_has_data = True
        if fund_has_data:
            ax4.set_xlabel("Year", fontsize=12, fontweight='bold'); ax4.set_ylabel("Calendar-Year Return (%)", fontsize=12, fontweight='bold')
            ax4.set_title("Individual Fund Calendar-Year Performance Comparison", fontsize=14, fontweight='bold')
            ax4.legend(loc='upper left', fontsize=9); ax4.grid(True, alpha=0.3); ax4.axhline(y=0, color='black', linestyle='-', linewidth=0.5)
            st.pyplot(fig4)
        else: st.warning("No individual fund calendar year data available.")

        # COMBINED CHARTS PDF DOWNLOAD
        st.subheader("📥 Download All Charts")
        buf_charts = io.BytesIO()
        with PdfPages(buf_charts) as pdf:
            fig_combined1 = plt.figure(figsize=(12, 6))
            fig_combined1.suptitle("Portfolio Allocation", fontsize=16, fontweight='bold', y=0.98)
            ax1 = fig_combined1.add_subplot(121)
            ax1.pie(equal_weights, labels=df['Fund Name'], autopct='%1.1f%%', startangle=90, colors=plt.cm.Paired.colors)
            ax1.set_title("Equal-Weighted Allocation")
            ax2 = fig_combined1.add_subplot(122)
            ax2.pie(opt_weights, labels=df['Fund Name'], autopct='%1.1f%%', startangle=90, colors=plt.cm.Paired.colors)
            ax2.set_title("Optimized Allocation")
            fig_combined1.tight_layout(rect=[0, 0, 1, 0.90])
            pdf.savefig(fig_combined1, bbox_inches='tight'); plt.close(fig_combined1)
            if has_data or has_benchmark: pdf.savefig(fig3, bbox_inches='tight')
            if fund_has_data: pdf.savefig(fig4, bbox_inches='tight')
        buf_charts.seek(0)
        st.download_button(label="📥 Download All Charts (Combined PDF)", data=buf_charts, file_name=f"Portfolio_Charts_{datetime.now().strftime('%Y%m%d')}.pdf", mime="application/pdf", use_container_width=True)

        st.subheader("📊 Portfolio Metrics Summary")
        col1, col2, col3, col4 = st.columns(4)
        with col1: st.metric("Risk-Adjusted Return", f"{opt_metrics['risk_adjusted']:.2f}", help="Return per unit of volatility")
        with col2: st.metric("Best Year", f"{opt_metrics['best_year']:.1f}%" if opt_metrics['best_year'] else "N/A")
        with col3: st.metric("Worst Year", f"{opt_metrics['worst_year']:.1f}%" if opt_metrics['worst_year'] else "N/A")
        with col4: st.metric("Consistency", f"{opt_metrics['consistency']:.0f}%", help="% of years with positive returns")

        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("← Back to Input", use_container_width=True, key="btn_back_input_analysis"):
                st.session_state.page = 'input'
                st.rerun()
        with col2:
            if st.button("Generate Report →", use_container_width=True, type="primary", key="btn_gen_report_analysis"):
                st.session_state.page = 'report'
                st.rerun()

    elif st.session_state.page == 'report':
        st.header("📄 Step 3: Portfolio Analysis Report")
        st.success("✅ Report generation is ready.")
        
        data = st.session_state.portfolio_data
        df = data['funds_df']
        year_returns_cols = data['year_returns']
        
        n = len(df)
        equal_weights = np.ones(n) / n
        if data['goal_type'] == "Reach a Target Sum ($)":
            target_return = calculate_required_cagr(data['target_value'], data['initial_investment'], data['monthly_contribution'], data['years'])
            target_amount = data['target_value']
        else:
            target_return = data['target_growth']
            target_amount = calculate_future_value(data['initial_investment'], data['monthly_contribution'], data['years'], target_return)
        
        opt_weights, opt_return, opt_vol = optimize_portfolio_max_return(df, data['risk_profile'])
        eq_metrics = calculate_portfolio_metrics(df, equal_weights, year_returns_cols)
        opt_metrics = calculate_portfolio_metrics(df, opt_weights, year_returns_cols)
        eq_amount = calculate_future_value(data['initial_investment'], data['monthly_contribution'], data['years'], eq_metrics['return'])
        opt_amount = calculate_future_value(data['initial_investment'], data['monthly_contribution'], data['years'], opt_metrics['return'])
        
        risk_thresholds = {"Conservative": 10.0, "Moderate": 15.0, "Growth": 20.0}
        max_vol_threshold = risk_thresholds[data['risk_profile']]
        
        doc = Document()
        title = doc.add_heading('Portfolio Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        if data.get('client_name'): doc.add_paragraph(f"Client Name: {data['client_name']}")
        doc.add_paragraph(f"Prepared For: {st.session_state.user_email}")
        doc.add_paragraph()
        
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(f"Goal Type: {data['goal_type']}")
        if data['goal_type'] == "Reach a Target Sum ($)": doc.add_paragraph(f"Target Amount: ${data['target_value']:,.0f}")
        else: doc.add_paragraph(f"Target Annual Growth: {data['target_growth']:.1f}%")
        doc.add_paragraph(f"Time Horizon: {data['years']} years")
        doc.add_paragraph("Analysis return basis: 1Y return where available, otherwise annualised 3Y and then annualised 5Y return.")
        doc.add_paragraph(f"Risk Profile: {data['risk_profile']} (Max Volatility: {max_vol_threshold}%)")
        doc.add_paragraph(f"Initial Investment: ${data['initial_investment']:,.0f}")
        doc.add_paragraph(f"Monthly Contribution: ${data['monthly_contribution']:,.0f}")
        doc.add_paragraph()
        
        doc.add_heading('Fund Details', level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Fund Name'; hdr_cells[1].text = 'Analysis Return (% p.a.)'; hdr_cells[2].text = 'Volatility (%)'
        analysis_returns = _analysis_return_series(df)
        hdr_cells[3].text = 'Mgmt Fee (%)'; hdr_cells[4].text = 'Equal Weight'; hdr_cells[5].text = 'Optimized Weight'
        for i, (_, row) in enumerate(df.iterrows()):
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Fund Name'])
            row_cells[1].text = f"{analysis_returns.iloc[i]:.2f}" if pd.notna(analysis_returns.iloc[i]) else "N/A"
            row_cells[2].text = f"{row['Volatility (%)']:.2f}" if pd.notna(row['Volatility (%)']) else "N/A"
            row_cells[3].text = f"{row['Mgmt Fee (%)']:.2f}" if pd.notna(row['Mgmt Fee (%)']) else "N/A"
            row_cells[4].text = f"{equal_weights[i]:.1%}"; row_cells[5].text = f"{opt_weights[i]:.1%}"
        doc.add_paragraph()
        
        doc.add_heading('Portfolio Performance Analysis', level=1)
        doc.add_paragraph(f"Equal-Weighted Return: {eq_metrics['return']:.1f}% p.a.")
        doc.add_paragraph(f"Optimized Return (Max): {opt_metrics['return']:.1f}% p.a.")
        doc.add_paragraph(f"Target Return: {target_return:.1f}% p.a.")
        doc.add_paragraph(f"Equal-Weighted Projected Amount: ${eq_amount:,.0f}")
        doc.add_paragraph(f"Optimized Projected Amount: ${opt_amount:,.0f}")
        doc.add_paragraph(f"Target Amount: ${target_amount:,.0f}")
        doc.add_paragraph()
        
        doc.add_heading('Risk Assessment', level=1)
        doc.add_paragraph(f"Risk Profile Threshold: {max_vol_threshold}% volatility")
        doc.add_paragraph(f"Equal-Weighted Volatility: {eq_metrics['volatility']:.1f}%")
        doc.add_paragraph(f"Optimized Volatility: {opt_metrics['volatility']:.1f}%")
        doc.add_paragraph(f"Risk-Adjusted Return: {opt_metrics['risk_adjusted']:.2f}")
        doc.add_paragraph()
        
        if opt_metrics['yearly_returns']:
            doc.add_heading('Historical Performance (Calendar Year Returns)', level=1)
            doc.add_paragraph(f"Best Year: {opt_metrics['best_year']:.1f}%")
            doc.add_paragraph(f"Worst Year: {opt_metrics['worst_year']:.1f}%")
            doc.add_paragraph(f"Average Annual Return: {opt_metrics['avg_yearly']:.1f}%")
            doc.add_paragraph(f"Consistency: {opt_metrics['consistency']:.0f}% of years with positive returns")
            doc.add_paragraph()
        
        doc.add_heading('Fee Impact Analysis', level=1)
        doc.add_paragraph(f"Average Management Fee: {opt_metrics['fee']:.2f}% p.a.")
        fee_impact_10yr = data['initial_investment'] * (1 - (1 - opt_metrics['fee']/100)**10)
        doc.add_paragraph(f"Estimated Fee Impact over 10 years: ${fee_impact_10yr:,.0f}")
        doc.add_paragraph()
        
        doc.add_heading('Recommendations (Based on Optimized Portfolio)', level=1)
        
        def find_monthly_for_target(target, initial, years, return_rate):
            r = return_rate / 100
            if r == 0: return max(0, (target - initial) / (years * 12))
            fv_initial = initial * ((1 + r) ** years)
            remaining = target - fv_initial
            if remaining <= 0: return 0
            return max(0, remaining * r / (12 * (((1 + r) ** years) - 1)))
        def find_initial_for_target(target, monthly, years, return_rate):
            r = return_rate / 100
            if r == 0: return max(0, target - (monthly * 12 * years))
            fv_monthly = monthly * 12 * (((1 + r) ** years) - 1) / r
            remaining = target - fv_monthly
            if remaining <= 0: return 0
            return max(0, remaining / ((1 + r) ** years))
        def find_years_for_target(target, initial, monthly, return_rate):
            for test_years in range(1, 100):
                if calculate_future_value(initial, monthly, test_years, return_rate) >= target: return test_years
            return 100

        req_monthly = find_monthly_for_target(target_amount, data['initial_investment'], data['years'], opt_metrics['return'])
        req_initial_zero_monthly = find_initial_for_target(target_amount, 0, data['years'], opt_metrics['return'])
        req_years = find_years_for_target(target_amount, data['initial_investment'], data['monthly_contribution'], opt_metrics['return'])

        diff_monthly = req_monthly - data['monthly_contribution']
        if diff_monthly < -1: doc.add_paragraph(f"Option 1: You can reduce your monthly contribution by ${abs(diff_monthly):,.0f} (New total: ${req_monthly:,.0f}/month).")
        elif diff_monthly > 1: doc.add_paragraph(f"Option 1: You need to increase your monthly contribution by ${diff_monthly:,.0f} (New total: ${req_monthly:,.0f}/month).")
        else: doc.add_paragraph("Option 1: Your current monthly contribution is exactly on track.")

        diff_initial = req_initial_zero_monthly - data['initial_investment']
        if diff_initial < -1: doc.add_paragraph(f"Option 2: If monthly contribution is $0, you can reduce initial investment by ${abs(diff_initial):,.0f} (New total: ${req_initial_zero_monthly:,.0f}).")
        elif diff_initial > 1: doc.add_paragraph(f"Option 2: If monthly contribution is $0, you need to increase initial investment by ${diff_initial:,.0f} (New total: ${req_initial_zero_monthly:,.0f}).")
        else: doc.add_paragraph("Option 2: If monthly contribution is $0, your current initial investment is exactly on track.")

        diff_years = req_years - data['years']
        if diff_years < 0: doc.add_paragraph(f"Option 3: You can reduce your time horizon by {abs(diff_years)} years (New total: {req_years} years).")
        elif diff_years > 0: doc.add_paragraph(f"Option 3: You need to increase your time horizon by {diff_years} years (New total: {req_years} years).")
        else: doc.add_paragraph("Option 3: Your current time horizon is exactly on track.")
        doc.add_paragraph()
        
        doc.add_heading('IMPORTANT DISCLAIMER & PROFESSIONAL GUIDANCE', level=1)
        doc.add_heading('1. Nature of This Analysis', level=2)
        doc.add_paragraph("This Portfolio Analysis Report is generated based on the specific inputs, assumptions, and historical data provided by the user via this tool. The calculations, optimized allocations, and projections (including recommendations to adjust contributions, initial capital, or time horizons) are mathematical models intended for educational and illustrative purposes only.")
        doc.add_heading('2. No Guarantee of Future Performance', level=2)
        doc.add_paragraph("The analysis relies on historical performance metrics (e.g., past returns, volatility) and stated fund objectives. Past performance is not indicative of future results. Market conditions, fund management changes, and economic factors can cause actual outcomes to differ materially from the projections shown in this report.")
        doc.add_heading('3. Not Professional Financial Advice', level=2)
        doc.add_paragraph("This tool and its outputs do not constitute personalized financial, investment, tax, or legal advice. The recommendations provided are generic and do not take into account your complete financial picture, liquidity needs, tax status, or other personal circumstances. You should not make any investment decisions solely based on this report.")
        doc.add_heading('4. Professional Consultation', level=2)
        doc.add_paragraph("While this tool provides a valuable high-level feasibility assessment, building a comprehensive, concrete wealth strategy requires a holistic review of your unique financial situation. If you would like to translate this analysis into an actionable, personalized investment strategy and execute it with professional oversight, please consult your investment adviser or reach out to:")
        p = doc.add_paragraph()
        p.add_run('Christopher Chew, CFP®, CFC®\n').bold = True
        doc.add_paragraph('Certified Financial Planner | Certified Business & Financial Coach')
        doc.add_paragraph('• Email: chrischew@acaplt.com')
        doc.add_paragraph('• Mobile/WhatsApp: +6012-213 9559')
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button(label="📄 Download Word Report (DOCX)", data=buffer, file_name=f"Portfolio_Analysis_{datetime.now().strftime('%Y%m%d')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        st.info("💡 **Note:** To convert the Word document to PDF, open it in Microsoft Word and use 'Save As' > 'PDF' format.")
        
        if st.button("← Back to Analysis", use_container_width=True, key="btn_back_analysis_report"):
            st.session_state.page = 'analysis'
            st.rerun()